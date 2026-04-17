from __future__ import annotations

import copy
import difflib
import json
import logging
import re
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

from app.services.llm_client import LlmClient
from app.services.storage import ConversationStorage, JsonStorage


logger = logging.getLogger(__name__)
BJ_TZ = timezone(timedelta(hours=8))


def _now_iso() -> str:
    return datetime.now(BJ_TZ).isoformat()


SECTION_TEMPLATES: list[dict[str, Any]] = [
    {
        "section_id": "overview",
        "title": "项目概述",
        "core_slots": ["solution", "stage_plan"],
        "writing_points": [
            "项目一句话定位与使命",
            "面向的人群与要解决的核心问题",
            "当前产品/服务形态与关键亮点",
            "团队当前所处阶段与近期节奏",
        ],
        "subheadings": ["项目定位与价值主张", "解决的问题与当前进展"],
        "frameworks": ["一句话价值主张", "使命/愿景/定位", "项目阶段雷达"],
    },
    {
        "section_id": "users",
        "title": "用户痛点与目标人群",
        "core_slots": ["target_user", "pain_point"],
        "writing_points": [
            "核心目标用户画像（身份、场景、规模估计）",
            "典型使用情境与痛点表现",
            "痛点背后的本质原因",
            "已观察到的用户访谈或行为证据",
        ],
        "subheadings": ["目标人群画像与场景", "核心痛点与成因"],
        "frameworks": ["用户画像 5W1H", "JTBD（用户想做的事）", "痛点分层（表层/结构/情感）"],
    },
    {
        "section_id": "solution",
        "title": "产品/服务方案",
        "core_slots": ["solution", "core_advantage"],
        "writing_points": [
            "产品或服务整体形态",
            "关键功能模块与服务流程",
            "用户从认知到使用的完整体验路径",
            "相较既有做法改善了什么",
        ],
        "subheadings": ["方案整体形态与核心功能", "使用流程与体验改进"],
        "frameworks": ["价值主张画布", "产品功能分层（核心/延伸/增强）", "体验地图"],
    },
    {
        "section_id": "business_model",
        "title": "商业模式与价值主张",
        "core_slots": ["business_model"],
        "writing_points": [
            "为谁创造价值，创造的是什么价值",
            "收费对象、收费方式与定价逻辑",
            "关键合作伙伴与收入来源组合",
            "价值交付与盈利的闭环",
        ],
        "subheadings": ["价值主张与收入模式", "合作伙伴与价值闭环"],
        "frameworks": ["商业模式画布九宫格", "收入模型拆解", "关键资源/合作伙伴矩阵"],
    },
    {
        "section_id": "market",
        "title": "市场与竞品分析",
        "core_slots": ["market_competition"],
        "writing_points": [
            "市场规模与增长趋势",
            "目标细分市场的进入切口",
            "主要竞品与替代方案画像",
            "自身与竞品的差异化定位",
        ],
        "subheadings": ["市场规模与趋势", "竞争格局与差异化定位"],
        "frameworks": ["TAM/SAM/SOM", "PEST 宏观分析", "波特五力", "竞争矩阵"],
    },
    {
        "section_id": "advantage",
        "title": "核心优势与竞争壁垒",
        "core_slots": ["core_advantage"],
        "writing_points": [
            "团队已具备的关键资源与能力",
            "产品或服务中难以被快速复制的点",
            "已有的数据、内容或生态积累",
            "长期可持续的壁垒来源",
        ],
        "subheadings": ["团队资源与能力壁垒", "生态积累与长期护城河"],
        "frameworks": ["SWOT", "护城河四象限（规模/网络/品牌/转换成本）", "壁垒时间线"],
    },
    {
        "section_id": "operations",
        "title": "运营与推广策略",
        "core_slots": ["operation_strategy"],
        "writing_points": [
            "目标用户在哪里出现，如何触达",
            "获客与激活路径，转化漏斗假设",
            "留存与复购的运营手段",
            "内容、活动或口碑机制",
        ],
        "subheadings": ["触达渠道与获客转化", "留存复购与内容口碑"],
        "frameworks": ["AARRR 海盗漏斗", "4P 营销组合", "渠道 ROI 矩阵"],
    },
    {
        "section_id": "finance",
        "title": "财务与融资计划",
        "core_slots": ["finance_logic"],
        "writing_points": [
            "启动投入、月度成本结构与假设",
            "收入模型、关键假设与敏感度",
            "盈亏平衡与资金缺口判断",
            "商业模式财务合理性评估与融资需求",
        ],
        "subheadings": ["成本结构与收入模型", "资金需求与财务合理性判断"],
        "frameworks": ["成本结构分层（固定/变动）", "损益三表框架", "盈亏平衡分析", "融资阶梯"],
    },
    {
        "section_id": "risk",
        "title": "当前风险与待验证假设",
        "core_slots": ["market_competition", "finance_logic"],
        "writing_points": [
            "业务假设中最关键、最未被验证的 2-3 条",
            "市场、竞争、政策、技术等维度的潜在风险",
            "对应的验证思路或缓解动作",
            "风险触发时的应对预案",
        ],
        "subheadings": ["关键未验证假设与外部风险", "验证路径与应对预案"],
        "frameworks": ["风险矩阵（概率×影响）", "假设-实验-指标", "应急预案分级"],
    },
    {
        "section_id": "roadmap",
        "title": "阶段目标与下一步行动",
        "core_slots": ["stage_plan"],
        "writing_points": [
            "近期 1-3 个月的阶段目标与里程碑",
            "中期 6-12 个月的目标设想",
            "近期需要落地的具体行动与责任分工",
            "验收标准与关键衡量指标",
        ],
        "subheadings": ["近期里程碑与行动", "中期路线图与关键指标"],
        "frameworks": ["里程碑甘特", "阶段 OKR", "北极星指标 + 输入指标"],
    },
]

MIN_SECTION_LENGTH = 420
PREFERRED_SECTION_LENGTH = 560

# 草稿态每章目标字数（has_material / stub）
DRAFT_TARGET_LEN = (300, 500)
DRAFT_STUB_TARGET_LEN = (260, 420)
# 升级为正式版（basic / full）时每章目标字数
UPGRADE_LEN = {
    "basic": {"material": (1200, 2000), "stub": (900, 1500)},
    "full": {"material": (2500, 4000), "stub": (1800, 2800)},
}

# 成熟度分组（3 组并发，按项目-用户、市场-价值、运营-落地拆分）
UPGRADE_GROUPS: list[list[str]] = [
    ["overview", "users", "solution"],
    ["market", "advantage", "business_model"],
    ["operations", "finance", "risk", "roadmap"],
]

# ── 成熟度评分常量 ────────────────────────────────────────────────
# 6 个骨架字段，每字段 10 分 = 60 分（相关字段二选一取高分）
_MATURITY_SKELETON_FIELDS: list[tuple[str, list[str]]] = [
    ("target_user", ["target_user"]),
    ("pain_point", ["pain_point"]),
    ("solution", ["solution"]),
    ("business_model", ["business_model"]),
    ("advantage_or_market", ["core_advantage", "market_competition"]),
    ("stage_or_ops", ["stage_plan", "operation_strategy"]),
]

# 信息具体性四档得分
_SPECIFICITY_SCORE = {
    "empty": 0,
    "vague": 4,
    "concrete": 7,
    "validated": 10,
}

_SPECIFICITY_LABEL = {
    "empty": "未提及",
    "vague": "模糊",
    "concrete": "具体",
    "validated": "已印证",
}

# 量化词、实体信号（用于判断字段信息是否"具体"）
_QUANT_PATTERN = re.compile(
    r"\d+|百分|千|万|亿|元|%|月|周|天|次|倍|第[一二三四五六七八九十]|"
    r"一二|三四|半年|季度"
)
_ENTITY_HINTS = (
    "高校", "中小学", "大学", "学生", "家长", "老师", "教师", "企业", "工厂", "医院",
    "APP", "小程序", "系统", "平台", "SaaS", "AI", "大模型", "算法", "OCR",
    "北京", "上海", "深圳", "广州", "杭州", "本地",
    "订阅", "会员", "按次", "抽成", "分成", "佣金", "授权", "广告",
)
_VAGUE_WORDS = ("一些", "大家", "某种", "大概", "可能", "也许", "不确定")

_MATURITY_TIER_LABEL = {
    "not_ready": "未就绪",
    "basic_ready": "基础就绪",
    "full_ready": "充分就绪",
}

CORE_SLOT_LABELS = {
    "target_user": "目标用户",
    "pain_point": "痛点",
    "solution": "方案",
    "core_advantage": "核心优势",
    "business_model": "商业模式",
    "market_competition": "市场/竞品",
    "operation_strategy": "运营策略",
    "finance_logic": "财务逻辑",
    "stage_plan": "阶段计划",
}

CORE_FIELD_ORDER = [
    "target_user",
    "pain_point",
    "solution",
    "core_advantage",
    "business_model",
    "market_competition",
    "operation_strategy",
    "finance_logic",
    "stage_plan",
]

FIELD_KEYWORDS = {
    "target_user": ("用户", "学生", "老师", "家长", "企业", "客户", "群体", "人群", "客群"),
    "pain_point": ("痛点", "问题", "麻烦", "低效", "困难", "需求", "困境"),
    "solution": ("方案", "系统", "平台", "产品", "服务", "功能", "工具", "解决"),
    "core_advantage": ("优势", "壁垒", "创新", "差异", "独特", "核心竞争力"),
    "business_model": ("收费", "盈利", "商业模式", "变现", "收入模式", "订阅", "佣金"),
    "market_competition": ("市场", "竞品", "竞争", "替代", "行业", "对手"),
    "operation_strategy": ("运营", "推广", "获客", "转化", "渠道", "增长", "传播"),
    "finance_logic": ("成本", "收入", "利润", "盈亏", "融资", "预算", "现金流", "回本"),
    "stage_plan": ("阶段", "计划", "里程碑", "本周", "下周", "执行", "落地", "路线"),
}

SECTION_MATERIAL_FIELDS: dict[str, list[str]] = {
    "overview": ["solution", "target_user", "pain_point", "stage_plan"],
    "users": ["target_user", "pain_point"],
    "solution": ["solution", "core_advantage"],
    "business_model": ["business_model", "finance_logic"],
    "market": ["market_competition"],
    "advantage": ["core_advantage", "solution"],
    "operations": ["operation_strategy", "target_user"],
    "finance": ["finance_logic", "business_model"],
    "risk": ["market_competition", "finance_logic", "target_user"],
    "roadmap": ["stage_plan", "solution"],
}

SECTION_PLACEHOLDER_HINT: dict[str, str] = {
    "overview": "项目整体定位与阶段进展",
    "users": "目标人群与核心痛点",
    "solution": "产品/服务方案形态",
    "business_model": "商业模式与收入结构",
    "market": "目标市场与主要竞品",
    "advantage": "核心优势与竞争壁垒",
    "operations": "运营与推广策略",
    "finance": "财务测算与融资计划",
    "risk": "关键风险与待验证假设",
    "roadmap": "阶段目标与下一步行动",
}


class BusinessPlanStorage:
    def __init__(self, root: Path) -> None:
        self.root = root
        self.root.mkdir(parents=True, exist_ok=True)
        self.index_file = self.root / "index.json"
        if not self.index_file.exists():
            self.index_file.write_text("[]", encoding="utf-8")

    def _load_index(self) -> list[dict[str, Any]]:
        try:
            return json.loads(self.index_file.read_text(encoding="utf-8"))
        except Exception:
            return []

    def _save_index(self, rows: list[dict[str, Any]]) -> None:
        self.index_file.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")

    def _plan_dir(self, project_id: str, conversation_id: str) -> Path:
        target = self.root / project_id / conversation_id
        target.mkdir(parents=True, exist_ok=True)
        return target

    def _plan_file(self, project_id: str, conversation_id: str, plan_id: str) -> Path:
        return self._plan_dir(project_id, conversation_id) / f"{plan_id}.json"

    def load(self, plan_id: str) -> dict[str, Any] | None:
        for row in self._load_index():
            if row.get("plan_id") != plan_id:
                continue
            target = self._plan_file(
                str(row.get("project_id") or ""),
                str(row.get("conversation_id") or ""),
                plan_id,
            )
            if not target.exists():
                continue
            try:
                return json.loads(target.read_text(encoding="utf-8"))
            except Exception:
                return None
        return None

    def load_latest(self, project_id: str, conversation_id: str) -> dict[str, Any] | None:
        plan_dir = self._plan_dir(project_id, conversation_id)
        files = sorted(plan_dir.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
        for file in files:
            try:
                return json.loads(file.read_text(encoding="utf-8"))
            except Exception:
                continue
        return None

    def list_by_project(self, project_id: str) -> list[dict[str, Any]]:
        project_id = str(project_id or "").strip()
        if not project_id:
            return []
        rows: list[dict[str, Any]] = []
        for row in self._load_index():
            if str(row.get("project_id") or "") != project_id:
                continue
            plan_id = str(row.get("plan_id") or "")
            conv_id = str(row.get("conversation_id") or "")
            target = self._plan_file(project_id, conv_id, plan_id)
            if not target.exists():
                continue
            try:
                data = json.loads(target.read_text(encoding="utf-8"))
            except Exception:
                continue
            rows.append({
                "plan_id": plan_id,
                "project_id": project_id,
                "conversation_id": conv_id,
                "project_name": data.get("project_name") or "",
                "one_liner": data.get("one_liner") or "",
                "version_tier": data.get("version_tier") or "",
                "updated_at": data.get("updated_at") or row.get("updated_at") or "",
                "readiness_score": (data.get("readiness") or {}).get("score"),
                "readiness_tier": (data.get("readiness") or {}).get("tier"),
            })
        rows.sort(key=lambda r: str(r.get("updated_at") or ""), reverse=True)
        return rows

    def save(self, plan: dict[str, Any]) -> dict[str, Any]:
        project_id = str(plan.get("project_id") or "").strip()
        conversation_id = str(plan.get("conversation_id") or "").strip()
        plan_id = str(plan.get("plan_id") or "").strip()
        if not project_id or not conversation_id or not plan_id:
            raise ValueError("business plan missing ids")
        target = self._plan_file(project_id, conversation_id, plan_id)
        target.write_text(json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")

        rows = self._load_index()
        meta = {
            "plan_id": plan_id,
            "project_id": project_id,
            "conversation_id": conversation_id,
            "updated_at": plan.get("updated_at") or _now_iso(),
        }
        found = False
        for idx, row in enumerate(rows):
            if row.get("plan_id") == plan_id:
                rows[idx] = {**row, **meta}
                found = True
                break
        if not found:
            rows.append(meta)
        self._save_index(rows)
        return plan


class BusinessPlanService:
    def __init__(
        self,
        storage: BusinessPlanStorage,
        json_store: JsonStorage,
        conv_store: ConversationStorage,
        llm: LlmClient | None = None,
    ) -> None:
        self.storage = storage
        self.json_store = json_store
        self.conv_store = conv_store
        self.llm = llm

    def get_latest(self, project_id: str, conversation_id: str) -> dict[str, Any] | None:
        return self.storage.load_latest(project_id, conversation_id)

    def get_plan(self, plan_id: str) -> dict[str, Any] | None:
        return self.storage.load(plan_id)

    def get_readiness(self, project_id: str, conversation_id: str) -> dict[str, Any]:
        conv = self.conv_store.get(project_id, conversation_id) or {}
        messages = list(conv.get("messages") or [])
        joined_text = "\n".join(str(msg.get("content") or "") for msg in messages)
        latest_trace = self._latest_trace(messages)
        exploration_state = latest_trace.get("exploration_state") if isinstance(latest_trace, dict) else {}
        filled = dict((exploration_state or {}).get("filled_slots") or {})

        slots: dict[str, bool] = {}
        for key, labels in FIELD_KEYWORDS.items():
            matched = any(token in joined_text for token in labels)
            if key == "target_user" and filled.get("target_user"):
                matched = True
            if key == "pain_point" and filled.get("pain_point"):
                matched = True
            if key == "solution" and filled.get("solution"):
                matched = True
            if key == "business_model" and filled.get("business_model"):
                matched = True
            slots[key] = matched

        core_keys = [
            "target_user",
            "pain_point",
            "solution",
            "business_model",
            "market_competition",
            "finance_logic",
            "stage_plan",
        ]
        filled_core = [key for key in core_keys if slots.get(key)]
        missing_core = [CORE_SLOT_LABELS[key] for key in core_keys if not slots.get(key)]

        fields = self._extract_core_fields(messages, latest_trace)
        raw_materials = self._harvest_raw_materials(messages)
        maturity = self._compute_maturity_score(
            fields=fields,
            raw_materials=raw_materials,
            latest_trace=latest_trace,
        )

        tier = maturity["tier"]
        ready = tier in {"basic_ready", "full_ready"} or len(filled_core) >= 4
        if tier == "not_ready":
            ready = False

        gap_questions = [item.get("suggestion") for item in maturity.get("next_gap", []) if item.get("suggestion")]
        suggested_questions = [q for q in gap_questions if q][:4] or [
            f"请补充一下项目的{label}。" for label in missing_core[:3]
        ]

        return {
            "ready": ready,
            "filled_core_slots": filled_core,
            "filled_core_count": len(filled_core),
            "missing_core_slots": missing_core,
            "suggested_questions": suggested_questions,
            "slot_map": slots,
            "message_count": len(messages),
            "maturity_score": maturity["score"],
            "maturity_tier": maturity["tier"],
            "maturity_tier_label": _MATURITY_TIER_LABEL.get(maturity["tier"], "未就绪"),
            "maturity_breakdown": maturity["breakdown"],
            "maturity_next_gap": maturity["next_gap"],
            "maturity_field_levels": maturity["field_levels"],
        }

    def generate_plan(
        self,
        *,
        project_id: str,
        conversation_id: str,
        student_id: str = "",
        allow_low_confidence: bool = False,
    ) -> dict[str, Any]:
        readiness = self.get_readiness(project_id, conversation_id)
        tier = str(readiness.get("maturity_tier") or "")
        # tier=not_ready 必须显式强制才能生成；其他 tier 正常走草稿流程
        if tier == "not_ready" and not allow_low_confidence:
            return {"status": "needs_more_info", "readiness": readiness, "plan": None}
        if not readiness.get("ready") and not allow_low_confidence and not tier:
            return {"status": "needs_more_info", "readiness": readiness, "plan": None}

        conv = self.conv_store.get(project_id, conversation_id)
        if not conv:
            return {"status": "not_found", "detail": "conversation not found", "plan": None}

        current = self.storage.load_latest(project_id, conversation_id)
        draft = self._build_draft(conv=conv, project_id=project_id, student_id=student_id, current=current)
        if current:
            revisions = self._build_revisions(current.get("sections") or [], draft.get("sections") or [], conv)
            draft["pending_revisions"] = revisions
            draft["revision_badge_count"] = len(revisions)
        else:
            draft["pending_revisions"] = []
            draft["revision_badge_count"] = 0

        saved = self._save_plan(draft, previous=current)
        return {"status": "ok", "plan": saved, "readiness": readiness}

    def refresh_plan(self, plan_id: str) -> dict[str, Any]:
        current = self.storage.load(plan_id)
        if not current:
            return {"status": "not_found", "plan": None}
        return self.generate_plan(
            project_id=str(current.get("project_id") or ""),
            conversation_id=str(current.get("conversation_id") or ""),
            student_id=str(current.get("student_id") or ""),
            allow_low_confidence=True,
        )

    def update_section(
        self,
        plan_id: str,
        section_id: str,
        *,
        content: str,
        field_map: dict[str, Any] | None = None,
        display_title: str | None = None,
    ) -> dict[str, Any] | None:
        plan = self.storage.load(plan_id)
        if not plan:
            return None

        next_sections: list[dict[str, Any]] = []
        for section in list(plan.get("sections") or []):
            row = copy.deepcopy(section)
            if row.get("section_id") == section_id:
                row["user_edit"] = content
                row["content"] = content
                if field_map is not None:
                    row["field_map"] = field_map
                if display_title is not None:
                    row["display_title"] = display_title
                row["updated_at"] = _now_iso()
                row["missing_level"] = self._missing_level(row)
                row["status"] = self._status_from_missing(row["missing_level"])
            next_sections.append(row)

        previous = copy.deepcopy(plan)
        plan["sections"] = next_sections
        plan["status"] = "user_edited"
        plan["updated_at"] = _now_iso()
        return self._save_plan(plan, previous=previous)

    def accept_revision(self, plan_id: str, revision_id: str) -> dict[str, Any] | None:
        plan = self.storage.load(plan_id)
        if not plan:
            return None
        revisions = list(plan.get("pending_revisions") or [])
        revision = next((row for row in revisions if row.get("revision_id") == revision_id), None)
        if not revision:
            return None

        previous = copy.deepcopy(plan)
        next_sections: list[dict[str, Any]] = []
        for section in list(plan.get("sections") or []):
            row = copy.deepcopy(section)
            if row.get("section_id") == revision.get("section_id"):
                row["ai_draft"] = revision.get("new_content") or row.get("ai_draft") or ""
                row["content"] = row["ai_draft"]
                row["user_edit"] = ""
                row["field_map"] = revision.get("candidate_field_map") or row.get("field_map") or {}
                row["missing_points"] = revision.get("candidate_missing_points") or row.get("missing_points") or []
                row["missing_level"] = revision.get("candidate_missing_level") or self._missing_level(row)
                row["status"] = self._status_from_missing(row["missing_level"])
                row["revision_status"] = "clean"
                row["updated_at"] = _now_iso()
            next_sections.append(row)

        plan["sections"] = next_sections
        plan["pending_revisions"] = [row for row in revisions if row.get("revision_id") != revision_id]
        plan["revision_badge_count"] = len(plan["pending_revisions"])
        plan["status"] = "synced"
        plan["updated_at"] = _now_iso()
        return self._save_plan(plan, previous=previous)

    def reject_revision(self, plan_id: str, revision_id: str) -> dict[str, Any] | None:
        plan = self.storage.load(plan_id)
        if not plan:
            return None
        previous = copy.deepcopy(plan)
        plan["pending_revisions"] = [
            row for row in list(plan.get("pending_revisions") or [])
            if row.get("revision_id") != revision_id
        ]
        plan["revision_badge_count"] = len(plan["pending_revisions"])
        plan["updated_at"] = _now_iso()
        return self._save_plan(plan, previous=previous)

    def export_stub(self, plan_id: str, export_mode: str, cover_info: dict[str, Any] | None = None) -> dict[str, Any]:
        return self.export_plan(
            plan_id=plan_id,
            export_mode=export_mode,
            export_format="docx",
            cover_info=cover_info,
        )

    def export_plan(
        self,
        *,
        plan_id: str,
        export_mode: str = "clean_final",
        export_format: str = "docx",
        cover_info: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found"}
        merged_cover = {**(plan.get("cover_info") or {}), **(cover_info or {})}

        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt, Cm, RGBColor  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
            from docx.oxml.ns import qn  # type: ignore
            from docx.oxml import OxmlElement  # type: ignore
        except Exception as exc:
            logger.warning("python-docx not available: %s", exc)
            return {
                "status": "unavailable",
                "message": "服务器未安装 python-docx，无法导出 docx。",
            }

        exports_dir = self.storage.root.parent / "exports"
        exports_dir.mkdir(parents=True, exist_ok=True)

        project_name = merged_cover.get("project_name") or plan.get("title") or "商业计划书"
        student_or_team = merged_cover.get("student_or_team") or ""
        course_or_class = merged_cover.get("course_or_class") or ""
        teacher_name = merged_cover.get("teacher_name") or ""
        cover_date = merged_cover.get("date") or _now_iso()[:10]

        ACCENT = RGBColor(0x1E, 0x3A, 0x5F)      # 深蓝 主色
        ACCENT_SOFT = RGBColor(0x45, 0x5A, 0x7D)  # 次级
        TEXT_MAIN = RGBColor(0x22, 0x26, 0x2E)
        TEXT_MUTED = RGBColor(0x80, 0x80, 0x80)
        STUB_COLOR = RGBColor(0x99, 0x99, 0x99)
        TABLE_HEADER_FILL = "E8EEF7"

        try:
            document = Document()
            sec = document.sections[0]
            sec.top_margin = Cm(2.2)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin = Cm(2.4)
            sec.right_margin = Cm(2.4)

            style = document.styles["Normal"]
            style.font.name = "微软雅黑"
            style.font.size = Pt(11)
            rpr = style.element.get_or_add_rPr()
            rFonts_el = rpr.find(qn("w:rFonts"))
            if rFonts_el is None:
                rFonts_el = OxmlElement("w:rFonts")
                rpr.append(rFonts_el)
            rFonts_el.set(qn("w:eastAsia"), "微软雅黑")
            try:
                style.paragraph_format.line_spacing = 1.5
                style.paragraph_format.space_after = Pt(6)
            except Exception:
                pass
        except Exception as exc:
            logger.exception("docx init failed: %s", exc)
            return {
                "status": "error",
                "export_format": export_format,
                "message": f"docx 初始化失败：{type(exc).__name__}: {exc}",
            }

        try:
            def _safe_set_eastasia(run_obj, font_name: str) -> None:
                try:
                    rpr_ = run_obj._element.get_or_add_rPr()
                    rf = rpr_.find(qn("w:rFonts"))
                    if rf is None:
                        rf = OxmlElement("w:rFonts")
                        rpr_.append(rf)
                    rf.set(qn("w:eastAsia"), font_name)
                except Exception:
                    pass

            def _set_cell_shading(cell, fill_hex: str) -> None:
                try:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), fill_hex)
                    tc_pr.append(shd)
                except Exception:
                    pass

            def _set_table_borders(table) -> None:
                try:
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    borders = OxmlElement("w:tblBorders")
                    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                        el = OxmlElement(f"w:{edge}")
                        el.set(qn("w:val"), "single")
                        el.set(qn("w:sz"), "4")
                        el.set(qn("w:color"), "8FA4C4")
                        borders.append(el)
                    tblPr.append(borders)
                except Exception:
                    pass

            def _parse_md_table(lines: list[str]) -> list[list[str]] | None:
                """接收形如 |..|..| 的连续行（含分隔行 |---|---|），返回二维字符串数组；失败返回 None。"""
                rows: list[list[str]] = []
                sep_seen = False
                for ln in lines:
                    s = ln.strip()
                    if not s.startswith("|"):
                        return None
                    cells = [c.strip() for c in s.strip("|").split("|")]
                    if all(re.match(r"^:?-{3,}:?$", c or "") for c in cells if c):
                        sep_seen = True
                        continue
                    rows.append(cells)
                if len(rows) < 1 or not sep_seen:
                    return None
                max_cols = max(len(r) for r in rows)
                for r in rows:
                    while len(r) < max_cols:
                        r.append("")
                return rows

            def _add_md_table(rows: list[list[str]]) -> None:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                table.autofit = True
                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        cell = table.cell(i, j)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        run_ = para.add_run(cell_text)
                        run_.font.size = Pt(10.5)
                        _safe_set_eastasia(run_, "微软雅黑")
                        if i == 0:
                            run_.bold = True
                            run_.font.color.rgb = ACCENT
                            _set_cell_shading(cell, TABLE_HEADER_FILL)
                _set_table_borders(table)

            def _add_section_content(md: str) -> None:
                """把章节 Markdown 解析为：标题 / 段落 / 列表 / 表格。"""
                raw_lines = md.splitlines()
                i = 0
                n = len(raw_lines)
                buffer_para: list[str] = []

                def _flush_para() -> None:
                    if not buffer_para:
                        return
                    text = " ".join(s.strip() for s in buffer_para if s.strip())
                    buffer_para.clear()
                    if not text:
                        return
                    p = document.add_paragraph()
                    r_ = p.add_run(text)
                    r_.font.size = Pt(11)
                    r_.font.color.rgb = TEXT_MAIN
                    _safe_set_eastasia(r_, "微软雅黑")

                while i < n:
                    ln = raw_lines[i]
                    st = ln.strip()

                    # 空行 → flush paragraph
                    if not st:
                        _flush_para()
                        i += 1
                        continue

                    # 表格识别：连续 | 开头行且包含分隔行
                    if st.startswith("|"):
                        tbl_lines = []
                        j = i
                        while j < n and raw_lines[j].strip().startswith("|"):
                            tbl_lines.append(raw_lines[j])
                            j += 1
                        rows = _parse_md_table(tbl_lines)
                        if rows:
                            _flush_para()
                            _add_md_table(rows)
                            i = j
                            continue

                    # 标题
                    if st.startswith("### "):
                        _flush_para()
                        h = document.add_heading(st[4:].strip(), level=3)
                        for sr in h.runs:
                            sr.font.name = "宋体"
                            sr.font.color.rgb = ACCENT_SOFT
                            _safe_set_eastasia(sr, "宋体")
                        i += 1
                        continue
                    if st.startswith("## "):
                        _flush_para()
                        h = document.add_heading(st[3:].strip(), level=2)
                        for sr in h.runs:
                            sr.font.name = "宋体"
                            sr.font.color.rgb = ACCENT
                            _safe_set_eastasia(sr, "宋体")
                        i += 1
                        continue

                    # 列表
                    if st.startswith("- ") or st.startswith("* "):
                        _flush_para()
                        while i < n and (raw_lines[i].strip().startswith("- ") or raw_lines[i].strip().startswith("* ")):
                            li = raw_lines[i].strip().lstrip("-*").strip()
                            if li:
                                try:
                                    p = document.add_paragraph(li, style="List Bullet")
                                except Exception:
                                    p = document.add_paragraph("• " + li)
                                for r_ in p.runs:
                                    _safe_set_eastasia(r_, "微软雅黑")
                                    r_.font.size = Pt(11)
                                    r_.font.color.rgb = TEXT_MAIN
                            i += 1
                        continue

                    # 普通段落
                    buffer_para.append(ln)
                    i += 1

                _flush_para()

            # ───── 封面 ─────
            for _ in range(3):
                document.add_paragraph()

            # 品牌带（小色块）
            accent_line = document.add_paragraph()
            accent_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            al_run = accent_line.add_run("商业计划书 / Business Plan")
            al_run.font.size = Pt(11)
            al_run.font.color.rgb = ACCENT
            al_run.bold = True
            _safe_set_eastasia(al_run, "微软雅黑")

            cover_title = document.add_paragraph()
            cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cover_title.add_run(str(project_name))
            run.font.name = "宋体"
            run.font.size = Pt(36)
            run.bold = True
            run.font.color.rgb = ACCENT
            _safe_set_eastasia(run, "宋体")

            kb = plan.get("knowledge_base") or {}
            one_liner = str(kb.get("one_liner") or "").strip()
            if one_liner:
                sub_title = document.add_paragraph()
                sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_run = sub_title.add_run(one_liner[:80])
                sub_run.font.size = Pt(13)
                sub_run.italic = True
                sub_run.font.color.rgb = ACCENT_SOFT
                _safe_set_eastasia(sub_run, "微软雅黑")

            for _ in range(5):
                document.add_paragraph()

            def _cover_row(label: str, value: str) -> None:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_ = p.add_run(f"{label}：{value or '＿＿＿＿＿＿'}")
                r_.font.size = Pt(12.5)
                r_.font.color.rgb = TEXT_MAIN
                _safe_set_eastasia(r_, "微软雅黑")

            _cover_row("团队 / 作者", student_or_team)
            _cover_row("课程 / 班级", course_or_class)
            _cover_row("指导老师", teacher_name)
            _cover_row("日期", cover_date)

            mat = plan.get("maturity") or {}
            if isinstance(mat, dict) and mat.get("score") is not None:
                mr_p = document.add_paragraph()
                mr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mr_run = mr_p.add_run(f"内容成熟度：{mat.get('score')}/100（{mat.get('tier_label') or mat.get('tier') or ''}）")
                mr_run.font.size = Pt(11)
                mr_run.font.color.rgb = TEXT_MUTED
                _safe_set_eastasia(mr_run, "微软雅黑")

            document.add_page_break()

            # ───── 目录 ─────
            toc_heading = document.add_heading("目  录", level=1)
            for sr in toc_heading.runs:
                sr.font.color.rgb = ACCENT
                sr.font.name = "宋体"
                _safe_set_eastasia(sr, "宋体")
            toc_para = document.add_paragraph()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            toc_para._element.append(fld_begin)

            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = ' TOC \\o "1-3" \\h \\z \\u '
            toc_para._element.append(instr)

            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            toc_para._element.append(fld_sep)

            toc_hint = OxmlElement("w:r")
            toc_hint_text = OxmlElement("w:t")
            toc_hint_text.text = "（请在 Word 中右键 → 更新域 以生成完整目录）"
            toc_hint.append(toc_hint_text)
            toc_para._element.append(toc_hint)

            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            toc_para._element.append(fld_end)

            document.add_page_break()

            # ───── 页眉 / 页脚 + 页码 ─────
            try:
                header = sec.header
                if header.paragraphs:
                    hp = header.paragraphs[0]
                else:
                    hp = header.add_paragraph()
                hp.text = ""
                hr = hp.add_run(str(project_name))
                hr.font.size = Pt(9)
                hr.font.color.rgb = TEXT_MUTED
                _safe_set_eastasia(hr, "微软雅黑")

                footer = sec.footer
                if footer.paragraphs:
                    fp = footer.paragraphs[0]
                else:
                    fp = footer.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fp.text = ""
                f_run = fp.add_run()
                fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
                f_run._r.append(fld1)
                instr2 = OxmlElement("w:instrText"); instr2.set(qn("xml:space"), "preserve")
                instr2.text = " PAGE "
                f_run._r.append(instr2)
                fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
                f_run._r.append(fld2)
                f_run.font.size = Pt(9)
                f_run.font.color.rgb = TEXT_MUTED
            except Exception as exc:
                logger.warning("header/footer set failed: %s", exc)

            # ───── 正文 ─────
            narrative_opening = ""
            for section in plan.get("sections") or []:
                if section.get("narrative_opening"):
                    narrative_opening = str(section.get("narrative_opening"))
                    break
            if narrative_opening:
                intro = document.add_paragraph()
                r_intro = intro.add_run(narrative_opening)
                r_intro.italic = True
                r_intro.font.size = Pt(11)
                r_intro.font.color.rgb = ACCENT_SOFT
                _safe_set_eastasia(r_intro, "微软雅黑")
                document.add_paragraph()

            for idx, section in enumerate(plan.get("sections") or []):
                heading_text = f"第 {idx + 1} 章  {section.get('display_title') or section.get('title') or ''}"
                heading = document.add_heading(heading_text, level=1)
                for h_run in heading.runs:
                    h_run.font.name = "宋体"
                    h_run.font.color.rgb = ACCENT
                    _safe_set_eastasia(h_run, "宋体")

                if section.get("is_ai_stub"):
                    note = document.add_paragraph()
                    nr = note.add_run("【AI 参考稿】本章基于项目知识库与行业通用框架生成，请团队校准。")
                    nr.italic = True
                    nr.font.size = Pt(10.5)
                    nr.font.color.rgb = STUB_COLOR
                    _safe_set_eastasia(nr, "微软雅黑")

                _add_section_content(str(section.get("content") or "").strip())

                if export_mode == "revision_marked":
                    for rev in plan.get("pending_revisions") or []:
                        if str(rev.get("section_id")) != str(section.get("section_id")):
                            continue
                        mark = document.add_paragraph()
                        mr = mark.add_run(f"【待审修订】{rev.get('summary') or ''}")
                        mr.italic = True
                        mr.font.size = Pt(10.5)
                        mr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                        _safe_set_eastasia(mr, "微软雅黑")

            # 文件命名
            ts = _now_iso().replace(":", "").replace("-", "")[:14]
            safe_name = re.sub(r"[\\/:*?\"<>|]+", "_", str(project_name))[:40] or "business_plan"
            docx_name = f"{safe_name}_{plan_id}_{ts}.docx"
            docx_path = exports_dir / docx_name
            document.save(docx_path)
        except Exception as exc:
            logger.exception("docx export failed: %s", exc)
            return {
                "status": "error",
                "export_format": export_format,
                "message": f"docx 生成失败：{type(exc).__name__}: {exc}",
            }

        if export_format == "pdf":
            try:
                from docx2pdf import convert  # type: ignore
                pdf_path = docx_path.with_suffix(".pdf")
                convert(str(docx_path), str(pdf_path))
                if pdf_path.exists():
                    return {
                        "status": "ok",
                        "export_mode": export_mode,
                        "export_format": "pdf",
                        "file_name": pdf_path.name,
                        "file_url": f"/api/business-plan/exports/{pdf_path.name}",
                        "cover_info": merged_cover,
                    }
            except Exception as exc:
                logger.warning("pdf export fallback to docx: %s", exc)
                return {
                    "status": "pdf_unavailable",
                    "export_mode": export_mode,
                    "export_format": "docx",
                    "file_name": docx_name,
                    "file_url": f"/api/business-plan/exports/{docx_name}",
                    "cover_info": merged_cover,
                    "message": (
                        "服务器未安装 docx2pdf 或缺少 MS Word/LibreOffice 运行环境，"
                        "已生成 docx 可直接下载（可用 Word 打开后另存为 PDF）。"
                    ),
                }

        return {
            "status": "ok",
            "export_mode": export_mode,
            "export_format": "docx",
            "file_name": docx_name,
            "file_url": f"/api/business-plan/exports/{docx_name}",
            "cover_info": merged_cover,
        }

    def _save_plan(self, plan: dict[str, Any], previous: dict[str, Any] | None) -> dict[str, Any]:
        plan_copy = copy.deepcopy(plan)
        if previous:
            snapshot = copy.deepcopy(previous)
            snapshot.pop("previous_version", None)
            plan_copy["previous_version"] = snapshot
            plan_copy["version"] = int(previous.get("version") or 1) + 1
        else:
            plan_copy.setdefault("version", 1)
            plan_copy.setdefault("previous_version", None)
        plan_copy.setdefault("created_at", _now_iso())
        plan_copy["updated_at"] = _now_iso()
        saved = self.storage.save(plan_copy)
        self._sync_project_meta(saved)
        return saved

    def _sync_project_meta(self, plan: dict[str, Any]) -> None:
        meta = {
            "plan_id": plan.get("plan_id"),
            "conversation_id": plan.get("conversation_id"),
            "title": plan.get("title") or "商业计划书",
            "status": plan.get("status") or "draft",
            "updated_at": plan.get("updated_at") or _now_iso(),
            "version": plan.get("version") or 1,
            "revision_badge_count": plan.get("revision_badge_count") or 0,
        }
        self.json_store.upsert_business_plan_meta(str(plan.get("project_id") or ""), meta)

    def _build_draft(
        self,
        *,
        conv: dict[str, Any],
        project_id: str,
        student_id: str,
        current: dict[str, Any] | None,
    ) -> dict[str, Any]:
        conversation_id = str(conv.get("conversation_id") or "")
        messages = list(conv.get("messages") or [])
        latest_trace = self._latest_trace(messages)
        latest_user = self._latest_user_message(messages)
        fields = self._extract_core_fields(messages, latest_trace)
        budget_hint = self._load_budget_hint(student_id)

        # 1) 素材拼盘（0 LLM）
        raw_materials = self._harvest_raw_materials(messages)
        # 2) KB Curator（1 LLM 或规则降级）
        kb = self._build_knowledge_base(raw_materials, budget_hint)
        # 3) 成熟度评分
        maturity = self._compute_maturity_score(
            fields=fields,
            raw_materials=raw_materials,
            latest_trace=latest_trace,
        )
        # 4) Draft Writer（1 LLM）；失败走规则兜底
        ai_sections = self._generate_draft_sections(
            kb=kb,
            fields=fields,
            latest_trace=latest_trace,
            budget_hint=budget_hint,
        ) or self._fallback_sections(fields, latest_trace, budget_hint)
        # 草稿态不做二次重写以控制耗时；长度不达标时沿用原内容即可
        ai_sections = self._ensure_draft_min_length(ai_sections, fields, latest_trace, budget_hint)

        if current:
            sections = self._merge_with_current_order(current.get("sections") or [], ai_sections)
            plan_id = str(current.get("plan_id") or "")
            created_at = str(current.get("created_at") or _now_iso())
        else:
            sections = ai_sections
            plan_id = str(uuid4())[:8]
            created_at = _now_iso()

        title = str(conv.get("title") or "").strip() or "商业计划书"
        return {
            "plan_id": plan_id,
            "project_id": project_id,
            "conversation_id": conversation_id,
            "student_id": student_id,
            "title": f"{title} · 商业计划书",
            "status": "draft" if not current else "synced",
            "version_tier": "draft",
            "created_at": created_at,
            "updated_at": _now_iso(),
            "sections": sections,
            "knowledge_base": kb,
            "maturity": maturity,
            "pending_revisions": [],
            "revision_badge_count": 0,
            "cover_info": {
                "project_name": title,
                "student_or_team": "",
                "course_or_class": "",
                "teacher_name": "",
                "date": _now_iso()[:10],
            },
            "source_summary": {
                "latest_user_message": latest_user[:200],
                "message_count": len(messages),
            },
        }

    def _ensure_draft_min_length(
        self,
        sections: list[dict[str, Any]],
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> list[dict[str, Any]]:
        """草稿态只保证极短内容有兜底文字，不触发二次 LLM 重写。"""
        for idx, template in enumerate(SECTION_TEMPLATES):
            if idx >= len(sections):
                break
            section = sections[idx]
            content = str(section.get("content") or "").strip()
            if self._chinese_length(content) >= 120:
                continue
            fallback = self._fallback_section_content(template, fields, latest_trace, budget_hint)
            section["content"] = fallback
            section["ai_draft"] = fallback
            section.setdefault("is_ai_stub", not self._section_has_material(template, fields, latest_trace, budget_hint))
            section["missing_level"] = self._missing_level(section)
            section["status"] = self._status_from_missing(section["missing_level"])
        return sections

    def _build_llm_context(
        self,
        *,
        messages: list[dict[str, Any]],
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> dict[str, str]:
        recent_messages: list[str] = []
        for msg in messages[-16:]:
            role = "学生" if msg.get("role") == "user" else "AI"
            content = str(msg.get("content") or "").strip()
            if content:
                recent_messages.append(f"{role}: {content[:360]}")
        recent_block = "\n".join(recent_messages) or "暂无"
        field_block = "\n".join(f"- {CORE_SLOT_LABELS[k]}: {v}" for k, v in fields.items() if v) or "- 暂无清晰核心字段"
        budget_block = json.dumps(budget_hint, ensure_ascii=False)[:1400] if budget_hint else "暂无"
        trace_block = json.dumps(
            {
                "next_task": latest_trace.get("next_task"),
                "diagnosis": latest_trace.get("diagnosis"),
                "kg_analysis": latest_trace.get("kg_analysis"),
                "planner": ((latest_trace.get("role_agents") or {}).get("planner") or {}),
            },
            ensure_ascii=False,
        )[:2200]
        return {
            "recent_block": recent_block,
            "field_block": field_block,
            "budget_block": budget_block,
            "trace_block": trace_block,
        }

    def _section_has_material(
        self,
        template: dict[str, Any],
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> bool:
        sid = template.get("section_id")
        required = SECTION_MATERIAL_FIELDS.get(str(sid), list(template.get("core_slots") or []))
        for key in required:
            if (fields.get(key) or "").strip():
                return True
        if sid == "finance":
            summary = (budget_hint or {}).get("summary") if isinstance(budget_hint, dict) else None
            if isinstance(summary, dict) and any(summary.get(k) for k in ("total_investment", "baseline_monthly_revenue", "funding_gap")):
                return True
        if sid == "roadmap":
            next_task = latest_trace.get("next_task") if isinstance(latest_trace.get("next_task"), dict) else {}
            if str((next_task or {}).get("title") or "").strip():
                return True
        if sid == "market":
            kg = latest_trace.get("kg_analysis") if isinstance(latest_trace.get("kg_analysis"), dict) else {}
            if str((kg or {}).get("insight") or "").strip():
                return True
        if sid == "risk":
            diag = latest_trace.get("diagnosis") if isinstance(latest_trace.get("diagnosis"), dict) else {}
            rules = diag.get("triggered_rules") if isinstance(diag.get("triggered_rules"), list) else []
            if rules:
                return True
        return False

    def _placeholder_content(self, template: dict[str, Any]) -> str:
        sid = str(template.get("section_id") or "")
        topic = SECTION_PLACEHOLDER_HINT.get(sid, template.get("title") or "本章")
        return f"_本章将在后续访谈中补齐关于「{topic}」的内容。_"

    def _build_section_spec_block(self, material_map: dict[str, bool] | None = None) -> str:
        rows: list[str] = []
        for idx, item in enumerate(SECTION_TEMPLATES):
            points = "；".join(item.get("writing_points") or [])
            subheads = "、".join(item.get("subheadings") or [])
            sid = str(item["section_id"])
            material_flag = ""
            if material_map is not None:
                has_material = bool(material_map.get(sid))
                material_flag = (
                    "\n   has_material: true （请按 500-800 字完整章节撰写）"
                    if has_material
                    else "\n   has_material: false （请只返回一句占位：\"本章将在后续访谈中补齐关于「"
                    + SECTION_PLACEHOLDER_HINT.get(sid, item["title"]) + "」的内容。\"）"
                )
            rows.append(
                f"{idx + 1}. {sid} · {item['title']}\n"
                f"   写作要点: {points or '基于对话内容自由组织'}\n"
                f"   建议小标题: {subheads or '可自拟'}"
                f"{material_flag}"
            )
        return "\n".join(rows)

    def _generate_sections_with_llm(
        self,
        *,
        messages: list[dict[str, Any]],
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> list[dict[str, Any]] | None:
        if not self.llm or not self.llm.enabled:
            return None

        ctx = self._build_llm_context(
            messages=messages,
            fields=fields,
            latest_trace=latest_trace,
            budget_hint=budget_hint,
        )
        material_map = {
            str(t["section_id"]): self._section_has_material(t, fields, latest_trace, budget_hint)
            for t in SECTION_TEMPLATES
        }
        spec_block = self._build_section_spec_block(material_map)

        system_prompt = (
            "你是一位资深商业计划书撰写顾问，长期为创业团队撰写正式的商业计划书。"
            "你的任务是根据提供的材料，写出一份 10 章的正式计划书文本，读起来要像顾问团队交付的成熟文档，而不是对话整理。\n\n"
            "【硬性写作规范】\n"
            "1. 视角与语气：始终使用第三人称书面语（以 \"项目\"、\"团队\" 为主语），严禁出现 \"学生说\"、\"用户提到\"、\"在对话中\"、\"你提到\" 等元叙述，也严禁把学生的原话作为引号引用。\n"
            "2. 加工而非复述：把提供的 \"思路原材料\" 当作线索，用顾问视角重新组织成有论证逻辑的段落；必要时可引入通用商业分析框架（用户画像、价值主张画布、商业模式画布、波特五力、PEST、AARRR 漏斗、竞争矩阵等）提升专业度。\n"
            "3. 气质自信克制：默认语气应该是自信但克制的顾问结论；不要反复使用 \"尚需验证\"、\"需要补充\"、\"暂未明确\"、\"目前还不够具体\" 这类口头禅。如果确实需要指出下一步，整章最多只在最后一段用 1-2 句话点到即止。\n"
            "4. 章节粒度：has_material=true 的章节必须写成 500-800 字、4-6 段、至少 2-3 个 \"### 小标题\" 的 Markdown 正文，段与段之间要有论点-论证-推论的逻辑推进。\n"
            "5. 占位章节：has_material=false 的章节不要强行扩写；content 只输出给定的那一句占位文本，confidence 设为 0.2 左右即可。\n"
            "6. 财务与融资计划章节：必须结合提供的预算 / 财务线索，给出对商业模式财务合理性的明确判断（例如是否具备盈亏平衡潜力、资金缺口是否合理、关键敏感度在哪里），不要只列数字。\n"
            "7. 引入框架或行业常识时，用自然的行文融入，不要让它像教科书罗列。\n\n"
            "【返回格式】\n"
            "返回一个 JSON 对象，必须包含 sections 数组。"
            "每个元素包含：section_id, title, content, missing_points, field_map, confidence。"
            "可选填 narrative_opening：80-160 字的引言段，用顾问口吻概述项目定位与文档脉络，放在阅读模式最上方。"
            "不要在 JSON 外输出任何解释文字。"
        )
        user_prompt = (
            "【章节结构与每章写作要求】\n"
            + spec_block
            + "\n\n【思路原材料 — 仅供顾问转化参考，禁止原样复用、禁止在正文中提及 \"学生 / 用户 / 对话\"】\n"
            + ctx["field_block"]
            + "\n\n【最近对话节选 — 同上，只做转化，不要复述】\n"
            + ctx["recent_block"]
            + "\n\n【结构化分析线索 — 任务 / 诊断 / 知识图谱 / 规划】\n"
            + ctx["trace_block"]
            + "\n\n【预算 / 财务线索】\n"
            + ctx["budget_block"]
            + "\n\n请现在以资深顾问的口吻产出完整的计划书 JSON。有素材的章节写成正式长文；占位章节只返回一句占位文本。"
        )
        raw = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.35)
        sections = raw.get("sections") if isinstance(raw, dict) else None
        if not isinstance(sections, list):
            return None
        narrative_opening = ""
        if isinstance(raw, dict):
            narrative_opening = str(raw.get("narrative_opening") or "").strip()
        normalized: list[dict[str, Any]] = []
        for template in SECTION_TEMPLATES:
            sid = str(template["section_id"])
            row = next((item for item in sections if str(item.get("section_id")) == sid), None) or {}
            has_material = bool(material_map.get(sid))
            raw_content = str(row.get("content") or "").strip()
            content = raw_content if has_material and raw_content else (raw_content if has_material else self._placeholder_content(template))
            if not has_material and raw_content:
                content = self._placeholder_content(template)
            section = self._materialize_section(
                section_id=sid,
                title=str(row.get("title") or template["title"]),
                content=content,
                field_map=row.get("field_map") if isinstance(row.get("field_map"), dict) else {},
                missing_points=row.get("missing_points") if isinstance(row.get("missing_points"), list) else [],
                confidence=float(row.get("confidence") or (0.65 if has_material else 0.2)),
                evidence_sources=["conversation", "agent_trace"],
            )
            section["has_material"] = has_material
            normalized.append(section)
        if narrative_opening and normalized:
            normalized[0]["narrative_opening"] = narrative_opening[:400]
        return normalized

    def _ensure_min_length(
        self,
        sections: list[dict[str, Any]],
        *,
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
        messages: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        if not sections:
            return sections
        for idx, template in enumerate(SECTION_TEMPLATES):
            if idx >= len(sections):
                break
            section = sections[idx]
            has_material = bool(section.get("has_material"))
            if not has_material:
                placeholder = self._placeholder_content(template)
                if str(section.get("content") or "").strip() != placeholder:
                    section["content"] = placeholder
                    section["ai_draft"] = placeholder
                    section["missing_level"] = self._missing_level(section)
                    section["status"] = self._status_from_missing(section["missing_level"])
                continue
            current_content = str(section.get("content") or "").strip()
            if self._chinese_length(current_content) >= MIN_SECTION_LENGTH:
                continue
            rewritten = self._rewrite_section(
                template=template,
                existing_content=current_content,
                fields=fields,
                latest_trace=latest_trace,
                budget_hint=budget_hint,
                messages=messages,
            )
            if rewritten and self._chinese_length(rewritten) > self._chinese_length(current_content):
                section["content"] = rewritten
                section["ai_draft"] = rewritten
                section["missing_level"] = self._missing_level(section)
                section["status"] = self._status_from_missing(section["missing_level"])
            elif self._chinese_length(current_content) < MIN_SECTION_LENGTH:
                padded = self._affirmative_fallback(template, fields, latest_trace, budget_hint)
                if self._chinese_length(padded) > self._chinese_length(current_content):
                    section["content"] = padded
                    section["ai_draft"] = padded
                    section["missing_level"] = self._missing_level(section)
                    section["status"] = self._status_from_missing(section["missing_level"])
        return sections

    def _rewrite_section(
        self,
        *,
        template: dict[str, Any],
        existing_content: str,
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
        messages: list[dict[str, Any]],
    ) -> str:
        if not self.llm or not self.llm.enabled:
            return ""
        ctx = self._build_llm_context(
            messages=messages,
            fields=fields,
            latest_trace=latest_trace,
            budget_hint=budget_hint,
        )
        points = "\n".join(f"- {item}" for item in (template.get("writing_points") or []))
        subheads = "、".join(template.get("subheadings") or []) or "可自拟 2-3 个三级小标题"
        system_prompt = (
            "你是一位资深商业计划书撰写顾问，现在只需要扩写一章正文。硬性要求：\n"
            "1. 使用第三人称书面语（主语以 \"项目 / 团队\" 为主），不得出现 \"学生 / 用户说 / 在对话中\" 等元叙述，也不得引用学生原话。\n"
            "2. 把线索当作原材料，用顾问视角重新组织；可以借助通用商业分析框架（用户画像、价值主张画布、商业模式画布、波特五力、PEST、AARRR 漏斗、竞争矩阵等）增强专业度。\n"
            "3. 默认语气自信克制，不要反复使用 \"尚需验证 / 需要补充 / 暂未明确\" 等口头禅；整章最多在最后一段用 1-2 句话点一次下一步方向。\n"
            "4. 输出纯 Markdown 正文，至少 500 个中文字符，4-6 段，含 2-3 个 \"### 小标题\"，段与段之间保持论点-论证-推论的逻辑推进。\n"
            "5. 不要输出任何 JSON、注释或元信息，只输出正文本身。"
        )
        user_prompt = (
            f"章节编号/标题：{template['section_id']} · {template['title']}\n"
            f"本章写作要点：\n{points}\n"
            f"建议小标题：{subheads}\n\n"
            "当前偏短的草稿（仅供参考，可大幅改写，但不要照抄）：\n"
            f"{existing_content or '（暂无现有草稿）'}\n\n"
            "【思路原材料 — 禁止原样复用】\n"
            f"{ctx['field_block']}\n\n"
            "【最近对话节选 — 只做转化，不要复述】\n"
            f"{ctx['recent_block']}\n\n"
            "【结构化分析线索】\n"
            f"{ctx['trace_block']}\n\n"
            "【预算 / 财务线索】\n"
            f"{ctx['budget_block']}\n\n"
            "请直接输出这一章的 Markdown 正文。"
        )
        try:
            return str(self.llm.chat_text(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.35) or "").strip()
        except Exception:
            try:
                raw = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.35)
                if isinstance(raw, dict):
                    return str(raw.get("content") or "").strip()
            except Exception:
                return ""
            return ""

    def _chinese_length(self, text: str) -> int:
        if not text:
            return 0
        return sum(1 for ch in text if ch.strip())

    def _fallback_sections(
        self,
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> list[dict[str, Any]]:
        """规则版草稿：LLM 不可用时兜底。"""
        results: list[dict[str, Any]] = []
        for template in SECTION_TEMPLATES:
            has_material = self._section_has_material(template, fields, latest_trace, budget_hint)
            field_map = self._field_map_for_section(template["section_id"], fields, budget_hint)
            missing_points = self._missing_points_for_section(template["core_slots"], fields)
            # 无论 has_material 与否，都用 _affirmative_fallback 产出结构化内容，
            # 无素材时置 is_ai_stub=true，前端给出参考稿徽章。
            content = self._affirmative_fallback(template, fields, latest_trace, budget_hint)
            confidence = max(0.45, 0.85 - 0.08 * len(missing_points)) if has_material else 0.35
            section = self._materialize_section(
                section_id=template["section_id"],
                title=template["title"],
                content=content,
                field_map=field_map,
                missing_points=missing_points,
                confidence=confidence,
                evidence_sources=["conversation", "agent_trace"],
            )
            section["has_material"] = has_material
            section["is_ai_stub"] = not has_material
            results.append(section)
        return results

    def _fallback_section_content(
        self,
        template: dict[str, Any],
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> str:
        """所有章节都返回结构化内容；无素材时由上层标记 is_ai_stub。"""
        return self._affirmative_fallback(template, fields, latest_trace, budget_hint)

    def _affirmative_fallback(
        self,
        template: dict[str, Any],
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> str:
        next_task = latest_trace.get("next_task") if isinstance(latest_trace.get("next_task"), dict) else {}
        planner = ((latest_trace.get("role_agents") or {}).get("planner") or {})
        planner_text = str(planner.get("analysis") or "").strip()
        diag = latest_trace.get("diagnosis") if isinstance(latest_trace.get("diagnosis"), dict) else {}
        risk_rules = diag.get("triggered_rules") if isinstance(diag.get("triggered_rules"), list) else []
        risk_names = [
            str(item.get("fallacy_label") or item.get("name") or "")
            for item in risk_rules[:3]
            if isinstance(item, dict) and (item.get("fallacy_label") or item.get("name"))
        ]
        budget_summary = (budget_hint or {}).get("summary") if isinstance(budget_hint, dict) else None

        def _distilled(key: str, default: str) -> str:
            value = (fields.get(key) or "").strip()
            if not value:
                return default
            value = re.sub(r"\s+", "", value)
            return value[:60] if len(value) > 60 else value

        target_user_key = _distilled("target_user", "核心目标人群")
        pain_key = _distilled("pain_point", "核心痛点")
        solution_key = _distilled("solution", "产品或服务方案")
        advantage_key = _distilled("core_advantage", "核心能力")
        business_model_key = _distilled("business_model", "商业模式")
        market_key = _distilled("market_competition", "目标市场")
        operation_key = _distilled("operation_strategy", "运营策略")
        finance_key = _distilled("finance_logic", "财务逻辑")
        stage_key = _distilled("stage_plan", str(next_task.get("title") or "近期阶段目标"))

        sid = template["section_id"]
        if sid == "overview":
            return self._compose_section(
                "项目定位",
                f"项目围绕 {solution_key} 展开，定位为一个面向 {target_user_key} 的创业方案，通过将成熟产品形态与行业一线场景相结合，"
                "构建出一条更契合实际使用路径的价值交付线。整体思路以价值主张画布为参照，把 \"用户想做的事 — 遇到的困扰 — 得到的收益\" 三层映射进项目目标。",
                "解决的问题",
                f"在现有供给环境中，{target_user_key} 长期面临 {pain_key} 的困扰，现有工具或服务要么覆盖不全、要么体验断层。"
                f"项目以 {solution_key} 为切入口，通过在关键环节做重构，打通从需求识别到结果交付的闭环。",
                "当前进展",
                f"团队当前已经推进至 {stage_key} 阶段，具备了可以与真实用户对齐预期的骨架。"
                "后续 1-3 个月将继续打磨核心体验，并以计划书中的章节结构作为迭代节奏的参照。",
                "未来脉络",
                "本文的后续章节将沿 \"用户 → 方案 → 商业模式 → 市场 → 优势 → 运营 → 财务 → 风险 → 路线\" 的脉络展开，"
                "并在相应章节引入行业通用分析框架，用顾问视角做出判断与建议。",
            )
        if sid == "users":
            return self._compose_section(
                "目标人群画像",
                f"项目的核心目标人群可归结为 {target_user_key}，从身份标签、使用场景和决策链条三个维度看，他们共享一组相似的行为模式与价值偏好。"
                "借助用户画像的分层方法，可以进一步拆出 1-2 个首发 persona，围绕其典型日程、工具栈和关键决策节点展开设计。",
                "核心痛点",
                f"在目前的供给结构下，这一人群最显著的痛点集中在 {pain_key}。这类痛点既表现为单次任务的效率损失，也会在长期使用中沉淀为 \"现有工具不够顺手\" 的心理摩擦，"
                "对愿付意愿、推荐意愿都会产生持续影响。",
                "痛点成因",
                "从供给侧看，痛点背后是信息不对称、流程冗余与体验割裂的叠加；从需求侧看，用户对同一任务的隐含期望正在升级，"
                "而现有方案未能及时跟上节奏。这类结构性错配恰恰是创业项目切入的机会窗口。",
                "验证与下一步",
                "项目将继续用访谈、问卷与小规模试用三种方式交叉验证上述假设，重点关注愿付意愿与任务频次两个硬指标，并据此决定 persona 的收敛方向。",
            )
        if sid == "solution":
            return self._compose_section(
                "方案总览",
                f"项目当前的产品/服务方案聚焦于 {solution_key}，通过把核心场景拆解成 \"感知 — 决策 — 执行 — 回顾\" 四个环节，"
                "围绕关键节点提供专门的工具与内容支撑，形成一个完整的体验闭环。",
                "核心功能模块",
                "整个产品由三层结构构成：第一层是核心任务流，保证用户能够顺畅完成关键操作；第二层是智能辅助与内容服务，用于降低认知负担、"
                "提升效率；第三层是数据沉淀与反馈机制，确保产品能够随使用不断进化。",
                "体验路径",
                "典型使用路径参考 AARRR 漏斗进行设计：Acquisition 通过内容和案例触达 → Activation 在首次使用中提供 \"小胜利\" → Retention 依靠周期性价值点维系"
                "→ Revenue 通过订阅或按次付费实现 → Referral 借助社区和口碑扩散。",
                "差异化改进",
                f"项目在 {advantage_key} 方向上做出了具体设计，让方案在效率、专业度或场景契合度上与现有做法形成可感知的差距，"
                "从而构成首批用户选择我们的关键理由。",
            )
        if sid == "business_model":
            return self._compose_section(
                "价值主张",
                f"在商业模式画布的维度下，项目为 {target_user_key} 输出的核心价值是 \"以更低摩擦解决 {pain_key}\"，"
                "这既是单次使用的体验改善，也是长期效率提升的累积收益。",
                "收入组合",
                f"项目规划的收入结构以 {business_model_key} 为主轴，辅以增值服务、会员订阅、企业侧打包等多元路径，"
                "以兼顾覆盖广度与客单价深度。不同收入点对留存、复购和现金流的影响各异，团队会根据首批数据选择主力组合。",
                "关键资源与伙伴",
                "关键资源包括产品能力、领域内容与核心运营人员；关键伙伴则覆盖内容/数据供应方、分发渠道与行业客户。"
                "合作关系按照互惠原则设计，避免单向依赖。",
                "盈亏闭环",
                "完整的商业闭环沿 \"交付价值 → 形成习惯 → 产生付费 → 反哺迭代\" 展开，关键在于留存率与付费转化率是否能稳定在行业基线之上，"
                "这两项指标将是判断模式可持续性的硬标准。",
            )
        if sid == "market":
            return self._compose_section(
                "市场规模与趋势",
                f"以 {market_key} 所在赛道为基准，行业整体仍处在结构升级期，需求侧的多样化与供给侧的能力缺口同时扩大，为新进入者留出了操作空间。"
                "可以通过权威行业报告 + 一手访谈的交叉方式量化规模，这里先用定性判断给出框架。",
                "细分切口",
                "项目优先从一个相对清晰的细分切口进入，该切口用户规模可控、画像相对一致，"
                "便于在早期集中资源验证假设；之后按 \"相邻人群 → 相邻场景 → 相邻地域\" 的顺序逐步扩展。",
                "竞争格局",
                "借助波特五力框架简化看：潜在进入者存在但门槛在升高，替代品主要是用户自研/拼装方案，供方与买方议价力相对分散，"
                "行业内现有竞争者之间存在功能重叠但服务深度不足。项目的机会点正是卡在 \"通用平台做不细 + 小众工具做不全\" 的空档。",
                "差异化定位",
                f"项目的差异化主要来自 {advantage_key} 与对细分场景的深度理解，以此形成 \"更贴近场景、更懂用户\" 的品牌感知，"
                "避免陷入与综合平台的正面竞争。",
            )
        if sid == "advantage":
            return self._compose_section(
                "资源与能力",
                f"团队目前沉淀的关键能力包括 {advantage_key}，这些能力在产品侧提供了具体落地路径，在认知侧提供了对行业的深层理解，"
                "构成了 \"敢做、能做、能做好\" 的基础。",
                "难以复制的环节",
                "项目中有若干环节难以被短时间复制：首先是围绕细分场景积累的方法论；其次是首批种子用户带来的反馈回路；"
                "最后是运营 / 产品 / 内容协同形成的节奏感，这种软性组织能力通常是后来者最难复制的部分。",
                "积累与生态",
                "随着项目推进，数据、案例、内容、社区等资产会逐步沉淀，并与产品能力互相增强，形成 \"数据 × 体验 × 生态\" 的组合式壁垒，"
                "这类壁垒会随时间持续加深。",
                "长期护城河",
                "从长远看，护城河的关键不是单点领先，而是多维叠加后的 \"替换成本\"：一旦用户把项目嵌入自己的日常工作流，"
                "换到竞品的决策成本会显著上升，这正是项目愿意在早期深耕的原因。",
            )
        if sid == "operations":
            return self._compose_section(
                "触达与渠道",
                f"项目的触达策略以 {operation_key} 为主线，在资源有限的早期阶段优先选择 ROI 明确、节奏快的 2-3 条渠道，"
                "通过小步试错沉淀可复用的运营资产，再据此决定加码或收缩。",
                "获客与转化",
                "在获客环节使用 AARRR 漏斗进行拆解：围绕 \"认知 → 点击 → 落地页 → 注册 → 首次关键动作\" 建立数据看板，"
                "针对每一级流失点制定对应的文案、流程优化与客服触达方案，保证转化曲线有稳定的改进方向。",
                "留存与复购",
                "留存围绕核心任务的周期性使用展开，用成就感反馈、个性化推送与社区归属感三条杠杆共同支撑；在条件成熟后引入会员或订阅，"
                "把偶发使用收敛为稳定付费关系。",
                "内容与口碑",
                "内容与口碑是长期压低获客成本的重要手段：项目会围绕目标人群真实痛点输出深度内容，借助早期用户故事在社区内形成 \"自来水\" 式传播，"
                "从而与付费投放形成互补。",
            )
        if sid == "finance":
            sm = budget_summary if isinstance(budget_summary, dict) else {}
            total_investment = sm.get("total_investment", 0)
            baseline_revenue = sm.get("baseline_monthly_revenue", 0)
            funding_gap = sm.get("funding_gap", 0)
            finance_block = (
                f"以现有预算测算，启动与早期运营的总投入约 {total_investment} 元，基准情境下月收入约 {baseline_revenue} 元，"
                f"短期资金缺口约 {funding_gap} 元。"
                if sm
                else "项目当前还未形成完整的量化测算，这里先从结构上给出顾问视角的财务判断，后续补齐细节。"
            )
            return self._compose_section(
                "成本结构",
                "项目成本按 \"产品研发 + 基本运营 + 获客投入\" 三部分建立基线。早期研发占比较高，"
                "进入稳定运营后运营与获客的比重会上升；财务模型会对这三项分别设定敏感度区间。",
                "收入模型",
                f"收入模型基于商业模式章节的 {business_model_key}，按不同增长速率与客单价假设分别测算 6/12/24 个月收入曲线，"
                "并标注每条曲线背后的关键假设，便于在后续迭代中快速调整。",
                "资金需求与盈亏",
                finance_block + "在此基础上进一步测算不同阶段的资金缺口、现金流安全边际，以及盈亏平衡所需的用户数或营收规模。",
                "财务合理性判断",
                f"从整体判断看，项目当前的财务逻辑（{finance_key}）具备阶段合理性：早期不追求短期盈利，而是以可控投入验证留存、转化与付费三项关键指标；"
                "一旦这三项稳定在行业基线之上，再加大获客投入是理性的扩张路径。若指标长期不达标，更合适的做法是调整收入结构而非继续融资续命。",
            )
        if sid == "risk":
            risk_text = "、".join(risk_names) if risk_names else "愿付意愿、获客成本、留存曲线"
            return self._compose_section(
                "关键未验证假设",
                "从顾问视角看，项目最关键的未验证假设集中在三点：目标用户是否具备稳定的愿付意愿、核心功能是否带来体验层级跃迁、"
                "获客成本是否能被控制在长期可负担的区间。这三点共同决定商业模式底座的稳固程度。",
                "外部与结构性风险",
                f"在更大的环境中，项目需要持续关注 {risk_text} 等维度；此外行业监管、数据合规、关键伙伴变化等外部变量也属于需定期复盘的风险条目。",
                "验证与缓解策略",
                "团队采用小步快跑的实验策略：以尽量低的成本设计可证伪的验证，围绕关键假设快速积累真实数据；"
                "同时在资源分配上保持必要冗余，避免在结论未清晰前将全部筹码押在单一方向。",
                "应对预案",
                "一旦关键假设被证伪或外部风险显著升级，预案包括：重新评估价值主张、调整收入结构、"
                "收缩或切换目标人群。这些动作不是退却，而是顾问视角下对不同情景的理性切换。",
            )
        if sid == "roadmap":
            planner_line = f"结合规划助手的提示：{planner_text[:120]}。" if planner_text else ""
            return self._compose_section(
                "近期阶段目标",
                f"项目近期 1-3 个月围绕 {stage_key} 推进，目标是完成一次从想法到可用方案的阶段性跃迁，"
                "形成一个可在真实场景中交付给目标用户的最小可行版本。",
                "中期设想",
                "展望 6-12 个月，项目会完成关键假设的系统验证，跑通第一个稳定的用户运营闭环，"
                "并为后续融资与扩张提供可度量的指标基础。每一个阶段节点都会基于数据做 \"加速 / 调整 / 收缩\" 的决策。",
                "近期具体行动",
                f"近期重点落地：完善核心功能、启动首批真实用户访谈与试用、搭建最小数据看板。{planner_line}"
                "团队内部同步建立责任分工与沟通节奏，以保证执行效率不在协作环节被稀释。",
                "关键指标",
                "衡量推进的关键指标包括：核心动作激活率与次周留存、真实用户反馈正向率、关键假设的验证进度。"
                "阶段回顾时对照这些指标，决定下一阶段的目标设定与资源投入方向。",
            )
        return self._compose_section(
            template["title"],
            "本章围绕当前可用的线索展开，按顾问视角给出结构化的分析与判断。",
            "主要思路",
            "结合本章的写作要点，项目会在此提供有逻辑的论述，而不是信息罗列。",
            "下一步",
            "后续几轮对话中将进一步补充背景、案例与数据，本章内容会据此持续迭代。",
        )

    def _compose_section(self, *segments: str) -> str:
        blocks: list[str] = []
        for idx in range(0, len(segments), 2):
            heading = segments[idx].strip()
            body = segments[idx + 1].strip() if idx + 1 < len(segments) else ""
            if heading:
                blocks.append(f"### {heading}")
            if body:
                blocks.append(body)
        return "\n\n".join(blocks)

    def _materialize_section(
        self,
        *,
        section_id: str,
        title: str,
        content: str,
        field_map: dict[str, Any],
        missing_points: list[str],
        confidence: float,
        evidence_sources: list[str],
    ) -> dict[str, Any]:
        section = {
            "section_id": section_id,
            "title": title,
            "display_title": title,
            "content": content.strip(),
            "ai_draft": content.strip(),
            "user_edit": "",
            "field_map": field_map,
            "missing_points": missing_points,
            "confidence": round(max(0.0, min(confidence, 1.0)), 2),
            "evidence_sources": evidence_sources,
            "revision_status": "clean",
            "is_custom": False,
            "updated_at": _now_iso(),
        }
        section["missing_level"] = self._missing_level(section)
        section["status"] = self._status_from_missing(section["missing_level"])
        return section

    def _field_map_for_section(self, section_id: str, fields: dict[str, str], budget_hint: dict[str, Any]) -> dict[str, Any]:
        section_fields: dict[str, Any] = {}
        for key in CORE_FIELD_ORDER:
            if fields.get(key):
                section_fields[key] = fields[key]
        if section_id == "finance" and budget_hint:
            summary = budget_hint.get("summary") or {}
            section_fields["budget_summary"] = {
                "total_investment": summary.get("total_investment", 0),
                "baseline_monthly_revenue": summary.get("baseline_monthly_revenue", 0),
                "funding_gap": summary.get("funding_gap", 0),
                "health_score": summary.get("health_score", 0),
            }
        return section_fields

    def _missing_points_for_section(self, core_slots: list[str], fields: dict[str, str]) -> list[str]:
        return [CORE_SLOT_LABELS[key] for key in core_slots if not fields.get(key)][:4]

    def _missing_level(self, section: dict[str, Any]) -> str:
        missing_count = len(list(section.get("missing_points") or []))
        content = str(section.get("content") or "").strip()
        content_len = self._chinese_length(content) if hasattr(self, "_chinese_length") else len(content)
        is_placeholder = content_len < 80 or content.startswith("_本章将在后续") or "将在后续访谈中补齐" in content
        if is_placeholder:
            return "critical"
        if missing_count == 0 and content_len >= 400:
            return "complete"
        if missing_count <= 1 and content_len >= 200:
            return "mostly_complete"
        if missing_count <= 3:
            return "partial"
        return "critical"

    def _status_from_missing(self, missing_level: str) -> str:
        return {
            "complete": "完整",
            "mostly_complete": "基本完整",
            "partial": "部分缺失",
            "critical": "严重缺失",
        }.get(missing_level, "部分缺失")

    def _merge_with_current_order(self, current_sections: list[dict[str, Any]], ai_sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
        ai_map = {str(row.get("section_id")): row for row in ai_sections}
        merged: list[dict[str, Any]] = []
        used: set[str] = set()
        for current in current_sections:
            sid = str(current.get("section_id") or "")
            candidate = copy.deepcopy(ai_map.get(sid, current))
            used.add(sid)
            candidate["display_title"] = current.get("display_title") or candidate.get("title")
            candidate["is_custom"] = bool(current.get("is_custom"))
            merged.append(candidate)
        for row in ai_sections:
            sid = str(row.get("section_id") or "")
            if sid not in used:
                merged.append(copy.deepcopy(row))
        return merged

    def _build_revisions(
        self,
        current_sections: list[dict[str, Any]],
        next_sections: list[dict[str, Any]],
        conv: dict[str, Any],
    ) -> list[dict[str, Any]]:
        current_map = {str(row.get("section_id")): row for row in current_sections}
        revisions: list[dict[str, Any]] = []
        source_hint = self._latest_user_message(list(conv.get("messages") or []))[:90]
        source_turn = sum(1 for msg in list(conv.get("messages") or []) if msg.get("role") == "user")
        for row in next_sections:
            sid = str(row.get("section_id") or "")
            prev = current_map.get(sid)
            if not prev:
                continue
            old_content = str(prev.get("content") or "")
            new_content = str(row.get("content") or "")
            old_norm = re.sub(r"\s+", " ", old_content).strip()
            new_norm = re.sub(r"\s+", " ", new_content).strip()
            if old_norm == new_norm:
                continue
            revisions.append(
                {
                    "revision_id": str(uuid4())[:8],
                    "section_id": sid,
                    "section_title": row.get("display_title") or row.get("title"),
                    "summary": f"{row.get('display_title') or row.get('title')} 有新的计划书修订建议",
                    "reason": "检测到新的对话信息，建议更新本章内容。",
                    "source_hint": f"基于第 {source_turn} 轮对话：{source_hint or '最新对话补充了新信息'}",
                    "old_content": old_content,
                    "new_content": new_content,
                    "candidate_field_map": row.get("field_map") or {},
                    "candidate_missing_points": row.get("missing_points") or [],
                    "candidate_missing_level": row.get("missing_level") or "partial",
                    "changes": self._build_inline_changes(old_content, new_content),
                    "created_at": _now_iso(),
                }
            )
        return revisions

    def _build_inline_changes(self, old_content: str, new_content: str) -> list[dict[str, str]]:
        diff = difflib.ndiff(old_content.splitlines() or [old_content], new_content.splitlines() or [new_content])
        rows: list[dict[str, str]] = []
        for item in diff:
            prefix = item[:2]
            text = item[2:]
            kind = "context"
            if prefix == "+ ":
                kind = "add"
            elif prefix == "- ":
                kind = "remove"
            if text.strip() or kind != "context":
                rows.append({"kind": kind, "text": text})
        return rows[:120]

    def _latest_trace(self, messages: list[dict[str, Any]]) -> dict[str, Any]:
        for msg in reversed(messages):
            trace = msg.get("agent_trace")
            if isinstance(trace, dict):
                return trace
        return {}

    def _latest_user_message(self, messages: list[dict[str, Any]]) -> str:
        for msg in reversed(messages):
            if msg.get("role") == "user":
                return str(msg.get("content") or "")
        return ""

    def _extract_core_fields(self, messages: list[dict[str, Any]], latest_trace: dict[str, Any]) -> dict[str, str]:
        recent_user_messages = [str(msg.get("content") or "") for msg in messages if msg.get("role") == "user"][-6:]
        all_user_text = "\n".join(recent_user_messages)
        latest_task = latest_trace.get("next_task") if isinstance(latest_trace.get("next_task"), dict) else {}
        kg = latest_trace.get("kg_analysis") if isinstance(latest_trace.get("kg_analysis"), dict) else {}
        diag = latest_trace.get("diagnosis") if isinstance(latest_trace.get("diagnosis"), dict) else {}
        planner = ((latest_trace.get("role_agents") or {}).get("planner") or {})

        fields = {
            "target_user": self._pick_sentence(all_user_text, FIELD_KEYWORDS["target_user"]),
            "pain_point": self._pick_sentence(all_user_text, FIELD_KEYWORDS["pain_point"]),
            "solution": self._pick_sentence(all_user_text, FIELD_KEYWORDS["solution"]),
            "core_advantage": self._pick_sentence(all_user_text, FIELD_KEYWORDS["core_advantage"]),
            "business_model": self._pick_sentence(all_user_text, FIELD_KEYWORDS["business_model"]),
            "market_competition": self._pick_sentence(all_user_text, FIELD_KEYWORDS["market_competition"]),
            "operation_strategy": self._pick_sentence(all_user_text, FIELD_KEYWORDS["operation_strategy"]),
            "finance_logic": self._pick_sentence(all_user_text, FIELD_KEYWORDS["finance_logic"]),
            "stage_plan": self._pick_sentence(all_user_text, FIELD_KEYWORDS["stage_plan"]),
        }
        if not fields["stage_plan"]:
            fields["stage_plan"] = str(latest_task.get("title") or "") or self._pick_sentence(str(planner.get("analysis") or ""), FIELD_KEYWORDS["stage_plan"])
        if not fields["finance_logic"]:
            fields["finance_logic"] = self._pick_sentence(str(diag.get("bottleneck") or ""), FIELD_KEYWORDS["finance_logic"])
        if not fields["market_competition"]:
            fields["market_competition"] = self._pick_sentence(str(kg.get("insight") or ""), FIELD_KEYWORDS["market_competition"])
        return {key: value for key, value in fields.items() if value}

    def _pick_sentence(self, text: str, keywords: tuple[str, ...]) -> str:
        if not text:
            return ""
        sentences = re.split(r"[。！？\n]", text)
        for sentence in sentences:
            item = sentence.strip()
            if len(item) < 6:
                continue
            if any(token in item for token in keywords):
                return item[:120]
        text = re.sub(r"\s+", " ", text).strip()
        return text[:120] if text and any(token in text for token in keywords) else ""

    # ══════════════════════════════════════════════════════════════
    #  Raw Material Harvest + KB Curator
    # ══════════════════════════════════════════════════════════════

    def _harvest_raw_materials(self, messages: list[dict[str, Any]]) -> dict[str, Any]:
        """聚合所有轮次的用户输入与多智能体产出。纯代码、0 次 LLM。

        目的是把散落在每轮 agent_trace 里的信息抽出来，分类、去重，作为 KB Curator 的输入。
        """
        user_turns: list[dict[str, Any]] = []
        assistant_turns: list[dict[str, Any]] = []
        planner_analyses: list[str] = []
        diagnosed_risks: list[dict[str, Any]] = []
        kg_entities: list[dict[str, Any]] = []
        kg_insights: list[str] = []
        proposed_tasks: list[dict[str, Any]] = []
        case_refs: list[dict[str, Any]] = []

        seen_user: set[str] = set()
        seen_planner: set[str] = set()
        seen_risk: set[str] = set()
        seen_entity: set[str] = set()
        seen_insight: set[str] = set()
        seen_task: set[str] = set()
        seen_case: set[str] = set()

        for idx, msg in enumerate(messages):
            role = msg.get("role")
            content = str(msg.get("content") or "").strip()
            if role == "user" and content:
                key = content[:60]
                if key not in seen_user:
                    seen_user.add(key)
                    user_turns.append({"round": idx + 1, "text": content})
            elif role in ("assistant", "ai") and content:
                assistant_turns.append({"round": idx + 1, "text": content[:800]})

            trace = msg.get("agent_trace")
            if not isinstance(trace, dict):
                continue

            planner = ((trace.get("role_agents") or {}).get("planner") or {})
            analysis = str(planner.get("analysis") or "").strip()
            if analysis:
                key = analysis[:60]
                if key not in seen_planner:
                    seen_planner.add(key)
                    planner_analyses.append(analysis)

            diag = trace.get("diagnosis") if isinstance(trace.get("diagnosis"), dict) else {}
            for rule in diag.get("triggered_rules") or []:
                if not isinstance(rule, dict):
                    continue
                name = str(rule.get("fallacy_label") or rule.get("name") or "").strip()
                if not name or name in seen_risk:
                    continue
                seen_risk.add(name)
                diagnosed_risks.append({
                    "name": name,
                    "severity": str(rule.get("severity") or "").strip(),
                    "evidence": str(rule.get("evidence") or "").strip()[:240],
                    "round": idx + 1,
                })

            kg = trace.get("kg_analysis") if isinstance(trace.get("kg_analysis"), dict) else {}
            for ent in kg.get("entities") or []:
                if not isinstance(ent, dict):
                    continue
                label = str(ent.get("name") or ent.get("label") or "").strip()
                if not label:
                    continue
                type_ = str(ent.get("type") or ent.get("category") or "Other").strip() or "Other"
                key = f"{type_}::{label}"
                if key in seen_entity:
                    continue
                seen_entity.add(key)
                kg_entities.append({"name": label, "type": type_})
            insight = str(kg.get("insight") or "").strip()
            if insight:
                key = insight[:60]
                if key not in seen_insight:
                    seen_insight.add(key)
                    kg_insights.append(insight)

            nt = trace.get("next_task") if isinstance(trace.get("next_task"), dict) else {}
            title = str(nt.get("title") or "").strip()
            if title and title not in seen_task:
                seen_task.add(title)
                proposed_tasks.append({
                    "title": title,
                    "description": str(nt.get("description") or "").strip()[:360],
                    "round": idx + 1,
                })

            cases = trace.get("case_knowledge") if isinstance(trace.get("case_knowledge"), dict) else {}
            for case in cases.get("matched") or cases.get("cases") or []:
                if not isinstance(case, dict):
                    continue
                cid = str(case.get("case_id") or case.get("id") or case.get("title") or "").strip()
                if not cid or cid in seen_case:
                    continue
                seen_case.add(cid)
                case_refs.append({
                    "case_id": cid,
                    "title": str(case.get("title") or "").strip()[:120],
                    "snippet": str(case.get("snippet") or case.get("summary") or "").strip()[:300],
                })

        return {
            "user_turns": user_turns,
            "assistant_turns": assistant_turns[-12:],
            "planner_analyses": planner_analyses[-8:],
            "diagnosed_risks": diagnosed_risks,
            "kg_entities": kg_entities,
            "kg_insights": kg_insights[-6:],
            "proposed_tasks": proposed_tasks,
            "case_refs": case_refs,
        }

    def _build_knowledge_base(
        self,
        raw_materials: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> dict[str, Any]:
        """KB Curator：1 次 LLM 调用，把素材拼盘蒸馏成结构化的项目知识库 JSON。

        若 LLM 不可用，退化为规则版 KB（保证后续逻辑可跑）。
        """
        fallback_kb = self._build_knowledge_base_fallback(raw_materials, budget_hint)
        if not self.llm or not self.llm.enabled:
            return fallback_kb

        materials_json = json.dumps(raw_materials, ensure_ascii=False)[:8000]
        budget_json = json.dumps(budget_hint or {}, ensure_ascii=False)[:1400]
        system_prompt = (
            "你是一位创业项目知识库编辑，专门把多轮用户对话与多智能体分析的原始素材，蒸馏成一份"
            "干净、结构化、可直接供撰写智能体使用的项目知识库 JSON。\n\n"
            "【硬性要求】\n"
            "1. 语气克制，只整理事实与明确的推论，不臆造数字；不确定的地方写入 open_questions。\n"
            "2. 去重、合并同义表达，避免把同一件事写三遍。\n"
            "3. 不得出现 '学生/用户说/在对话中' 等元叙述；全部用第三人称事实陈述。\n"
            "4. 严格返回 JSON 对象，schema 如下：\n"
            "{\n"
            "  \"project_core\": {\"one_liner\": \"\", \"mission\": \"\", \"stage\": \"\"},\n"
            "  \"user_insights\": {\n"
            "    \"target_personas\": [{\"label\": \"\", \"traits\": [], \"evidence_round\": 0}],\n"
            "    \"jobs_to_be_done\": [],\n"
            "    \"pain_points_ranked\": [{\"text\": \"\", \"severity\": \"\"}]\n"
            "  },\n"
            "  \"solution_design\": {\"value_props\": [], \"core_features\": [], \"differentiation\": []},\n"
            "  \"market_competitive\": {\"segments\": [], \"competitors_mentioned\": [], \"market_size_clues\": []},\n"
            "  \"business_economics\": {\"revenue_hypotheses\": [], \"cost_structure_hints\": [], \"budget_facts\": {}},\n"
            "  \"operations\": {\"channels\": [], \"funnel_assumptions\": []},\n"
            "  \"risks_identified\": [{\"name\": \"\", \"source_round\": 0, \"severity\": \"\"}],\n"
            "  \"agent_consensus\": [],\n"
            "  \"open_questions\": []\n"
            "}\n"
        )
        user_prompt = (
            "【素材拼盘（全部来源，禁止反向引用 '学生/对话'）】\n"
            f"{materials_json}\n\n"
            "【预算线索（可直接塞入 business_economics.budget_facts）】\n"
            f"{budget_json}\n\n"
            "请严格按 schema 输出 JSON。"
        )
        try:
            kb = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.2)
        except Exception as exc:
            logger.warning("KB Curator failed: %s", exc)
            return fallback_kb
        if not isinstance(kb, dict) or not kb:
            return fallback_kb
        # 补齐缺失字段，避免下游访问 KeyError
        for key, default in fallback_kb.items():
            kb.setdefault(key, default)
        return kb

    def _build_knowledge_base_fallback(
        self,
        raw_materials: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> dict[str, Any]:
        """规则版 KB：在 LLM 不可用时保证流程能跑通。"""
        risks = [
            {"name": r.get("name"), "source_round": r.get("round"), "severity": r.get("severity")}
            for r in (raw_materials.get("diagnosed_risks") or [])[:6]
        ]
        entities = raw_materials.get("kg_entities") or []
        budget_summary = (budget_hint or {}).get("summary") if isinstance(budget_hint, dict) else {}
        return {
            "project_core": {"one_liner": "", "mission": "", "stage": ""},
            "user_insights": {
                "target_personas": [],
                "jobs_to_be_done": [],
                "pain_points_ranked": [],
            },
            "solution_design": {"value_props": [], "core_features": [], "differentiation": []},
            "market_competitive": {
                "segments": [],
                "competitors_mentioned": [e.get("name") for e in entities if e.get("type") == "Competitor"][:5],
                "market_size_clues": [],
            },
            "business_economics": {
                "revenue_hypotheses": [],
                "cost_structure_hints": [],
                "budget_facts": budget_summary if isinstance(budget_summary, dict) else {},
            },
            "operations": {"channels": [], "funnel_assumptions": []},
            "risks_identified": risks,
            "agent_consensus": (raw_materials.get("planner_analyses") or [])[-3:],
            "open_questions": [t.get("title") for t in (raw_materials.get("proposed_tasks") or [])[:4]],
        }

    # ══════════════════════════════════════════════════════════════
    #  Maturity Score
    # ══════════════════════════════════════════════════════════════

    def _field_specificity(
        self,
        value: str,
        kg_entities: list[dict[str, Any]],
    ) -> str:
        """返回 empty / vague / concrete / validated 四档。"""
        text = (value or "").strip()
        if not text:
            return "empty"
        length = self._chinese_length(text)
        has_quant = bool(_QUANT_PATTERN.search(text))
        has_entity = any(hint in text for hint in _ENTITY_HINTS)
        has_vague = any(word in text for word in _VAGUE_WORDS)

        anchor_count = sum([has_quant, has_entity])
        if length >= 12 and anchor_count >= 2 and not (has_vague and anchor_count < 2):
            level = "concrete"
        elif length >= 8 and (has_quant or has_entity):
            level = "concrete"
        elif length >= 4:
            level = "vague"
        else:
            level = "vague"

        if level == "concrete":
            for ent in kg_entities:
                name = str(ent.get("name") or "")
                if name and name in text:
                    return "validated"
        return level

    def _compute_maturity_score(
        self,
        *,
        fields: dict[str, str],
        raw_materials: dict[str, Any],
        latest_trace: dict[str, Any],
    ) -> dict[str, Any]:
        """按 骨架60 + 智能体30 + 逻辑10 打分，返回 {score, tier, breakdown, next_gap, field_levels}。"""
        kg_entities = raw_materials.get("kg_entities") or []

        # A. 骨架完整度 (60)
        field_levels: dict[str, str] = {}
        skeleton_points = 0.0
        for group_name, candidate_keys in _MATURITY_SKELETON_FIELDS:
            best_level = "empty"
            best_key = candidate_keys[0]
            for key in candidate_keys:
                level = self._field_specificity(fields.get(key, ""), kg_entities)
                if _SPECIFICITY_SCORE[level] > _SPECIFICITY_SCORE[best_level]:
                    best_level = level
                    best_key = key
            field_levels[group_name] = best_level
            field_levels[f"{group_name}__key"] = best_key
            skeleton_points += _SPECIFICITY_SCORE[best_level]
        skeleton_score = min(round(skeleton_points), 60)

        # B. 多智能体信息密度 (30)
        entity_types = {str(e.get("type") or "") for e in kg_entities}
        relevant_types = {"User", "Problem", "Solution", "Competitor", "Resource", "Scenario"}
        type_coverage = len(entity_types & relevant_types)
        kg_score = min(type_coverage * 2, 10)

        risk_count = len(raw_materials.get("diagnosed_risks") or [])
        risk_score = min(risk_count * 3, 10)

        actionable = len(raw_materials.get("proposed_tasks") or []) * 2 + len(raw_materials.get("case_refs") or []) * 2
        actionable_score = min(actionable, 10)

        agent_score = kg_score + risk_score + actionable_score

        # C. 逻辑自洽 (10)
        coherence_score = 0
        pain_text = fields.get("pain_point", "") or ""
        solution_text = fields.get("solution", "") or ""
        business_text = fields.get("business_model", "") or ""
        if pain_text and solution_text:
            overlap = self._keyword_overlap(pain_text, solution_text)
            if overlap >= 2:
                coherence_score += 5
            elif overlap == 1:
                coherence_score += 3
        if solution_text and business_text:
            overlap = self._keyword_overlap(solution_text, business_text)
            if overlap >= 2:
                coherence_score += 5
            elif overlap == 1:
                coherence_score += 3

        total = skeleton_score + agent_score + coherence_score
        if total >= 65:
            tier = "full_ready"
        elif total >= 40:
            tier = "basic_ready"
        else:
            tier = "not_ready"

        next_gap = self._compute_next_gap(
            field_levels=field_levels,
            fields=fields,
            raw_materials=raw_materials,
            skeleton_score=skeleton_score,
            agent_score=agent_score,
            coherence_score=coherence_score,
        )

        return {
            "score": total,
            "tier": tier,
            "breakdown": {
                "skeleton": skeleton_score,
                "agent_density": agent_score,
                "coherence": coherence_score,
                "skeleton_max": 60,
                "agent_density_max": 30,
                "coherence_max": 10,
            },
            "next_gap": next_gap,
            "field_levels": {k: v for k, v in field_levels.items() if not k.endswith("__key")},
        }

    def _keyword_overlap(self, a: str, b: str) -> int:
        """粗略的关键词重合度：把中文按 2 字 n-gram 切分后求交集大小。"""
        def _grams(text: str) -> set[str]:
            clean = re.sub(r"[\s，。、,.:;:;!！?？]+", "", text)
            return {clean[i:i + 2] for i in range(max(len(clean) - 1, 0))}
        return len(_grams(a) & _grams(b))

    def _compute_next_gap(
        self,
        *,
        field_levels: dict[str, str],
        fields: dict[str, str],
        raw_materials: dict[str, Any],
        skeleton_score: int,
        agent_score: int,
        coherence_score: int,
    ) -> list[dict[str, Any]]:
        gaps: list[dict[str, Any]] = []

        # 字段层面的缺口（挑 3 个最弱的）
        field_scored: list[tuple[str, str]] = []
        for group_name, _candidates in _MATURITY_SKELETON_FIELDS:
            level = field_levels.get(group_name, "empty")
            field_scored.append((group_name, level))
        field_scored.sort(key=lambda x: _SPECIFICITY_SCORE[x[1]])

        readable = {
            "target_user": "目标用户",
            "pain_point": "核心痛点",
            "solution": "方案设计",
            "business_model": "商业模式",
            "advantage_or_market": "核心优势或市场",
            "stage_or_ops": "阶段计划或运营",
        }
        suggestion_map = {
            "target_user": "再具体描述一下目标人群的身份、场景与规模。",
            "pain_point": "聊一下这些用户最痛的 1-2 个场景或瞬间。",
            "solution": "详细讲一下产品/服务的形态、关键功能或使用流程。",
            "business_model": "说说怎么赚钱：付费方、定价逻辑、收入来源组合。",
            "advantage_or_market": "谈谈核心优势或者目标细分市场的情况。",
            "stage_or_ops": "说明当前所处阶段、下一步动作，或运营策略。",
        }
        for group_name, level in field_scored[:3]:
            if _SPECIFICITY_SCORE[level] >= 10:
                continue
            gaps.append({
                "dimension": "skeleton",
                "field": group_name,
                "field_label": readable.get(group_name, group_name),
                "current_level": level,
                "current_level_label": _SPECIFICITY_LABEL.get(level, level),
                "reason": f"{readable.get(group_name, group_name)} 当前档位为「{_SPECIFICITY_LABEL.get(level, level)}」",
                "suggestion": suggestion_map.get(group_name, f"补充{readable.get(group_name, group_name)}的更多细节。"),
            })

        # 智能体密度缺口
        if agent_score < 20:
            if not raw_materials.get("diagnosed_risks"):
                gaps.append({
                    "dimension": "agent_density",
                    "field": "risks",
                    "reason": "诊断引擎尚未捕获任何风险条目",
                    "suggestion": "聊一聊你最担心的 2-3 个业务风险或待验证假设。",
                })
            if not raw_materials.get("kg_entities"):
                gaps.append({
                    "dimension": "agent_density",
                    "field": "kg",
                    "reason": "知识图谱尚未抽取到关键实体",
                    "suggestion": "补充一个典型场景或具体案例，帮助识别关键用户、方案与竞品。",
                })
            if not raw_materials.get("proposed_tasks"):
                gaps.append({
                    "dimension": "agent_density",
                    "field": "tasks",
                    "reason": "还未沉淀出可落地的近期动作",
                    "suggestion": "说说你打算本月或下周落地的 1-2 件具体事情。",
                })

        # 逻辑自洽缺口
        if coherence_score < 5:
            gaps.append({
                "dimension": "coherence",
                "field": "pain_solution_loop",
                "reason": "痛点与方案之间的对应关系不够清晰",
                "suggestion": "具体说明方案中哪一个功能或服务对应解决哪一类痛点。",
            })

        return gaps[:5]

    # ══════════════════════════════════════════════════════════════
    #  Draft Writer (first-generation, 1 LLM call)
    # ══════════════════════════════════════════════════════════════

    def _build_section_spec_for_draft(
        self,
        material_map: dict[str, bool],
    ) -> str:
        rows: list[str] = []
        for idx, item in enumerate(SECTION_TEMPLATES):
            sid = str(item["section_id"])
            points = "；".join(item.get("writing_points") or [])
            frameworks = "、".join(item.get("frameworks") or []) or "按行业常识展开"
            has_material = material_map.get(sid, False)
            length_hint = (
                f"300-500 字短段落 + 3-5 个 Markdown 关键点 bullet" if has_material
                else f"260-420 字 AI 参考稿骨架 + 3-5 个 bullet（置 is_ai_stub=true）"
            )
            rows.append(
                f"{idx + 1}. {sid} · {item['title']}\n"
                f"   写作要点：{points}\n"
                f"   建议引入框架：{frameworks}\n"
                f"   has_material：{str(has_material).lower()}\n"
                f"   本章目标：{length_hint}"
            )
        return "\n".join(rows)

    def _generate_draft_sections(
        self,
        *,
        kb: dict[str, Any],
        fields: dict[str, str],
        latest_trace: dict[str, Any],
        budget_hint: dict[str, Any],
    ) -> list[dict[str, Any]] | None:
        """1 次 LLM 调用产出全 10 章的短草稿（骨架 + 关键点）。"""
        if not self.llm or not self.llm.enabled:
            return None

        material_map = {
            str(t["section_id"]): self._section_has_material(t, fields, latest_trace, budget_hint)
            for t in SECTION_TEMPLATES
        }
        spec_block = self._build_section_spec_for_draft(material_map)
        kb_json = json.dumps(kb, ensure_ascii=False)[:6000]
        budget_json = json.dumps(budget_hint or {}, ensure_ascii=False)[:1200]

        system_prompt = (
            "你是一位资深商业计划书撰写顾问。现在的任务是：基于提供的项目知识库（KB），"
            "一次性产出一份 10 章商业计划书的草稿（骨架版），让团队能快速通读、编辑与决定是否升级为正式版。\n\n"
            "【硬性写作规范】\n"
            "1. 视角与语气：始终使用第三人称书面语（主语以 '项目 / 团队' 为主），严禁出现 '学生 / 用户说 / 在对话中' 等元叙述，也严禁引用原话。\n"
            "2. 加工而非复述：把 KB 当成事实线索，用顾问视角组织成有论证逻辑的段落；适度引入每章对应的分析框架，但不要像教科书罗列。\n"
            "3. 语气自信克制：默认是顾问结论，不要反复 '尚需验证 / 需要补充 / 暂未明确' 等口头禅；整章最多在最后一段点一次下一步即可。\n"
            "4. 章节粒度：has_material=true 的章节写 300-500 字的 Markdown 正文（1-2 个 ### 小标题）+ 紧随其后的 3-5 个关键点 Markdown bullet。\n"
            "   has_material=false 的章节按对应框架写 260-420 字的 AI 参考稿骨架 + 3-5 个 bullet，需要在开头写一句 '以下为基于行业通用框架生成的参考稿，需团队校准。' 并在返回里置 is_ai_stub=true。\n"
            "5. 财务章节必须基于预算线索给出对商业模式财务合理性的明确判断（是否盈亏平衡潜力、资金缺口是否合理、敏感度所在），不要只列数字。\n\n"
            "【返回格式】\n"
            "返回 JSON 对象：{\"narrative_opening\": \"80-160 字的顾问引言\", \"sections\": [{\"section_id\", \"title\", \"content\", \"bullets\": [\"...\"], \"is_ai_stub\": bool, \"confidence\": 0.0-1.0, \"missing_points\": [...]}]}\n"
            "content 与 bullets 分开存储（content 是段落，bullets 是关键点数组）。\n"
            "不要在 JSON 外输出任何解释文字。"
        )
        user_prompt = (
            "【章节规格】\n" + spec_block
            + "\n\n【项目知识库（KB）— 所有事实的唯一来源】\n" + kb_json
            + "\n\n【预算线索】\n" + budget_json
            + "\n\n请现在一次性产出 10 章的草稿 JSON。"
        )
        try:
            raw = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.35)
        except Exception as exc:
            logger.warning("Draft Writer failed: %s", exc)
            return None

        sections = raw.get("sections") if isinstance(raw, dict) else None
        if not isinstance(sections, list):
            return None
        narrative_opening = ""
        if isinstance(raw, dict):
            narrative_opening = str(raw.get("narrative_opening") or "").strip()

        normalized: list[dict[str, Any]] = []
        for template in SECTION_TEMPLATES:
            sid = str(template["section_id"])
            row = next((item for item in sections if str(item.get("section_id")) == sid), None) or {}
            has_material = bool(material_map.get(sid))
            raw_content = str(row.get("content") or "").strip()
            bullets = row.get("bullets") if isinstance(row.get("bullets"), list) else []
            bullets = [str(b).strip() for b in bullets if str(b or "").strip()][:6]

            is_stub = bool(row.get("is_ai_stub")) or (not has_material)

            content_md = raw_content
            if bullets:
                bullets_md = "\n".join(f"- {b}" for b in bullets)
                content_md = (raw_content + "\n\n" + bullets_md).strip() if raw_content else bullets_md

            if not content_md:
                content_md = self._affirmative_fallback(template, fields, latest_trace, budget_hint)
                is_stub = not has_material

            section = self._materialize_section(
                section_id=sid,
                title=str(row.get("title") or template["title"]),
                content=content_md,
                field_map=row.get("field_map") if isinstance(row.get("field_map"), dict) else self._field_map_for_section(sid, fields, budget_hint),
                missing_points=row.get("missing_points") if isinstance(row.get("missing_points"), list) else self._missing_points_for_section(template["core_slots"], fields),
                confidence=float(row.get("confidence") or (0.55 if has_material else 0.3)),
                evidence_sources=["knowledge_base", "conversation", "agent_trace"],
            )
            section["has_material"] = has_material
            section["is_ai_stub"] = is_stub
            section["bullets"] = bullets
            normalized.append(section)

        if narrative_opening and normalized:
            normalized[0]["narrative_opening"] = narrative_opening[:400]
        return normalized

    # ══════════════════════════════════════════════════════════════
    #  Formal Upgrade (3-group concurrent, LLM)
    # ══════════════════════════════════════════════════════════════

    def upgrade_plan(
        self,
        plan_id: str,
        *,
        mode: str = "full",
    ) -> dict[str, Any]:
        """把草稿升级为正式版（full）或基础版（basic）。"""
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found", "plan": None}
        if mode not in UPGRADE_LEN:
            mode = "full"

        kb = plan.get("knowledge_base") or {}
        # 合并深化阶段学生补充的 chapter_addons 到 KB（仅为本次 LLM prompt 注入，不改 plan 存储）
        addons = kb.get("chapter_addons") or {}
        sections = list(plan.get("sections") or [])
        section_map = {str(s.get("section_id")): s for s in sections}
        requested_ids = [str(t["section_id"]) for t in SECTION_TEMPLATES]

        draft_map = {
            sid: {
                "title": section_map.get(sid, {}).get("display_title") or section_map.get(sid, {}).get("title") or "",
                "content": section_map.get(sid, {}).get("content") or "",
                "is_ai_stub": bool(section_map.get(sid, {}).get("is_ai_stub")),
                "has_material": bool(section_map.get(sid, {}).get("has_material")),
            }
            for sid in requested_ids
        }

        web_ctx = self._collect_web_context(requested_ids, kb)
        case_header, case_per_section = self._retrieve_case_fewshots(kb)

        upgraded_map: dict[str, dict[str, Any]] = {}
        outlines_map: dict[str, Any] = {}
        if self.llm and self.llm.enabled:
            # === 新：整章两步式并发 ===
            def _per_chapter_kb(sid: str) -> dict[str, Any]:
                extra = addons.get(sid)
                if not extra:
                    return kb
                kb2 = copy.deepcopy(kb)
                kb2.setdefault("chapter_focus", {})[sid] = (
                    extra if isinstance(extra, list) else [str(extra)]
                )
                return kb2

            with ThreadPoolExecutor(max_workers=5) as pool:
                futures = {
                    pool.submit(
                        self._upgrade_chapter_two_step,
                        sid, _per_chapter_kb(sid), draft_map, mode,
                        web_ctx.get(sid, ""),
                        case_header,
                        case_per_section.get(sid, ""),
                    ): sid
                    for sid in requested_ids
                }
                for fut, sid in futures.items():
                    try:
                        result = fut.result(timeout=260)
                    except Exception as exc:
                        logger.warning("two-step upgrade failed for %s: %s", sid, exc)
                        result = None
                    if result and result.get("content"):
                        upgraded_map[sid] = result
                        if result.get("outline"):
                            outlines_map[sid] = result["outline"]

            # === 兜底 1：对失败章节退回到 _upgrade_single ===
            still_missing = [sid for sid in requested_ids if sid not in upgraded_map]
            for sid in still_missing:
                try:
                    single = self._upgrade_single(
                        sid, _per_chapter_kb(sid), draft_map, mode,
                        web_ctx.get(sid, ""),
                        case_header, case_per_section.get(sid, ""),
                    )
                except Exception as exc:
                    logger.warning("_upgrade_single fallback failed for %s: %s", sid, exc)
                    continue
                if single:
                    upgraded_map[sid] = single

            # === 兜底 2：仍缺章的用 _upgrade_group 打包一次 ===
            still_missing = [sid for sid in requested_ids if sid not in upgraded_map]
            if still_missing:
                try:
                    group_result = self._upgrade_group(
                        still_missing, kb, draft_map, mode,
                        web_ctx, case_header, case_per_section,
                    )
                except Exception as exc:
                    logger.warning("_upgrade_group final fallback failed: %s", exc)
                    group_result = {}
                if isinstance(group_result, dict):
                    for k, v in group_result.items():
                        upgraded_map.setdefault(k, v)

        previous = copy.deepcopy(plan)
        next_sections: list[dict[str, Any]] = []
        revisions: list[dict[str, Any]] = []
        for section in sections:
            sid = str(section.get("section_id") or "")
            row = copy.deepcopy(section)
            upgraded = upgraded_map.get(sid)
            if upgraded and upgraded.get("content"):
                new_content = str(upgraded.get("content") or "").strip()
                if new_content and new_content != row.get("content"):
                    revisions.append({
                        "revision_id": str(uuid4())[:8],
                        "section_id": sid,
                        "section_title": row.get("display_title") or row.get("title"),
                        "summary": f"{row.get('display_title') or row.get('title')} 升级为正式版",
                        "reason": "基于项目知识库 + 行业参考资料对本章进行扩写与论证强化。",
                        "source_hint": f"升级模式：{mode}",
                        "old_content": row.get("content") or "",
                        "new_content": new_content,
                        "candidate_field_map": row.get("field_map") or {},
                        "candidate_missing_points": upgraded.get("missing_points") or row.get("missing_points") or [],
                        "candidate_missing_level": "complete",
                        "changes": self._build_inline_changes(row.get("content") or "", new_content),
                        "created_at": _now_iso(),
                    })
                    row["ai_draft"] = new_content
                    row["revision_status"] = "pending"
                    if "is_ai_stub" in upgraded:
                        row["is_ai_stub"] = bool(upgraded["is_ai_stub"])
                    if upgraded.get("outline"):
                        row["outline"] = upgraded["outline"]
            next_sections.append(row)

        plan_copy = copy.deepcopy(plan)
        plan_copy["sections"] = next_sections
        plan_copy["pending_revisions"] = (plan_copy.get("pending_revisions") or []) + revisions
        plan_copy["revision_badge_count"] = len(plan_copy["pending_revisions"])
        plan_copy["version_tier"] = mode
        plan_copy["status"] = "upgraded" if revisions else plan_copy.get("status", "draft")
        plan_copy["updated_at"] = _now_iso()
        plan_copy["upgrade_report"] = {
            "mode": mode,
            "requested": requested_ids,
            "success_ids": sorted(upgraded_map.keys()),
            "failed_ids": [sid for sid in requested_ids if sid not in upgraded_map],
            "timestamp": _now_iso(),
        }
        saved = self._save_plan(plan_copy, previous=previous)
        return {"status": "ok", "plan": saved}

    def _collect_web_context(
        self,
        group_ids: list[str],
        kb: dict[str, Any],
    ) -> dict[str, str]:
        """每章抓取 5-8 条相关行业资料并摘要为 800 字以内的参考文本。"""
        try:
            from app.services.web_search import web_search, format_for_llm
        except Exception as exc:
            logger.warning("web_search unavailable: %s", exc)
            return {}

        kw_map: dict[str, list[Any]] = {
            "overview":       [kb.get("project_name"), kb.get("one_liner"), "商业计划书范本", "项目定位"],
            "users":          [kb.get("target_user"), kb.get("pain_point"), "用户画像", "jobs-to-be-done"],
            "solution":       [kb.get("solution"), kb.get("project_name"), "产品功能设计", "MVP"],
            "market":         [kb.get("market_segment"), "市场规模 TAM SAM SOM", "行业趋势", "竞争格局"],
            "advantage":      [kb.get("core_advantage"), kb.get("project_name"), "竞争壁垒", "差异化"],
            "business_model": [kb.get("business_model"), "商业模式画布", "收入模型", "单位经济"],
            "operations":     [kb.get("operation_strategy"), "AARRR", "运营增长", "获客成本"],
            "finance":        [kb.get("project_name"), "财务预测", "LTV CAC 单位经济", "融资需求"],
            "risk":           [kb.get("project_name"), "创业 风险分析", "合规风险"],
            "roadmap":        [kb.get("project_name"), "产品路线图", "发展阶段规划", "里程碑"],
        }
        ctx: dict[str, str] = {}
        for sid in group_ids:
            kws = [str(k).strip() for k in kw_map.get(sid, []) if k]
            if not kws:
                continue
            query = " ".join(kws[:4])[:140]
            try:
                res = web_search(query, intent=sid, max_results=8)
                snippet = format_for_llm(res, max_chars=800)
                if snippet:
                    ctx[sid] = snippet
            except Exception as exc:
                logger.warning("web_search failed for %s: %s", sid, exc)
        return ctx

    # ── Few-shot 样本：从本地案例库挑 2 份同行业优秀计划书 ──
    _CASE_SECTION_MAP: dict[str, list[str]] = {
        "overview":       ["project_name", "solution"],
        "users":          ["target_users", "pain_points"],
        "solution":       ["solution", "innovation_points"],
        "market":         ["market_analysis"],
        "advantage":      ["innovation_points"],
        "business_model": ["business_model"],
        "operations":     ["execution_plan"],
        "finance":        ["business_model"],
        "risk":           ["risk_control"],
        "roadmap":        ["execution_plan"],
    }

    def _retrieve_case_fewshots(self, kb: dict[str, Any]) -> tuple[str, dict[str, str]]:
        """
        基于 KB 推断行业类别，取 2 份同类优秀案例：
        - 返回 (header_text, per_section_text)
        - header_text: 写入全局 prompt 顶部，告诉 LLM 风格参考哪些范本
        - per_section_text[sid]: 单章写作时可附的具体案例片段
        """
        try:
            from app.services.case_knowledge import infer_category, retrieve_cases_by_category
        except Exception as exc:
            logger.warning("case_knowledge unavailable: %s", exc)
            return "", {}

        hint_text = " ".join(
            str(v) for v in [
                kb.get("project_name"), kb.get("one_liner"), kb.get("market_segment"),
                kb.get("solution"), kb.get("core_advantage"),
            ] if v
        ) or "创新创业项目"
        try:
            category = infer_category(hint_text)
        except Exception:
            category = "科技创新"

        try:
            case_refs = retrieve_cases_by_category(category, limit=2)
        except Exception as exc:
            logger.warning("retrieve_cases failed: %s", exc)
            return "", {}
        if not case_refs:
            return "", {}

        structured_dir = self.storage.root.parent / "graph_seed" / "case_structured"
        loaded: list[dict[str, Any]] = []
        for ref in case_refs:
            cid = str(ref.get("case_id") or "")
            if not cid:
                continue
            p = structured_dir / f"{cid}.json"
            if not p.exists():
                continue
            try:
                data = json.loads(p.read_text(encoding="utf-8"))
                loaded.append(data)
            except Exception as exc:
                logger.warning("load case %s failed: %s", cid, exc)

        if not loaded:
            return "", {}

        header_lines = [
            f"【同行业优秀计划书范本（{category}，仅供写作风格、深度与结构参考，不得抄袭事实）】",
        ]
        for case in loaded:
            profile = case.get("project_profile") or {}
            header_lines.append(
                f"- 范本：{profile.get('project_name', '未命名')}"
                f"｜亮点：{'、'.join((profile.get('innovation_points') or [])[:2])}"
            )
        header = "\n".join(header_lines)

        per_section: dict[str, str] = {}
        for sid, fields in self._CASE_SECTION_MAP.items():
            blocks: list[str] = []
            for case in loaded:
                profile = case.get("project_profile") or {}
                case_name = profile.get("project_name", "范本")
                parts: list[str] = []
                for fld in fields:
                    val = profile.get(fld)
                    if isinstance(val, list) and val:
                        parts.append("；".join(str(x) for x in val[:4]))
                    elif isinstance(val, str) and val:
                        parts.append(val)
                if parts:
                    blocks.append(f"· {case_name}：{' / '.join(parts)[:500]}")
            if blocks:
                per_section[sid] = "\n".join(blocks)[:900]
        return header, per_section

    def _build_upgrade_spec(
        self,
        sid: str,
        draft_map: dict[str, dict[str, Any]],
        mode: str,
    ) -> tuple[str, str] | None:
        """返回 (spec_block, draft_block)。找不到模板返回 None。"""
        template = next((t for t in SECTION_TEMPLATES if t["section_id"] == sid), None)
        if not template:
            return None
        lens = UPGRADE_LEN.get(mode, UPGRADE_LEN["full"])
        draft = draft_map.get(sid, {})
        has_material = bool(draft.get("has_material"))
        len_range = lens["material"] if has_material else lens["stub"]
        points = "；".join(template.get("writing_points") or [])
        subheads = "、".join(template.get("subheadings") or [])
        frameworks = "、".join(template.get("frameworks") or [])
        spec = (
            f"- {sid} · {template['title']}\n"
            f"  写作要点：{points}\n"
            f"  小标题：{subheads}\n"
            f"  必须引入框架：{frameworks}\n"
            f"  has_material：{str(has_material).lower()}\n"
            f"  目标字数：{len_range[0]}-{len_range[1]} 字，4-6 段，至少 2 个 ### 小标题"
        )
        draft_block = f"### [{sid} 现有草稿]\n" + (draft.get("content") or "（暂无草稿）")
        return spec, draft_block

    def _upgrade_system_prompt(self, mode: str) -> str:
        return (
            "你是一位资深商业计划书撰写顾问。现在要把指定几章的短草稿升级为正式长版本。\n\n"
            "【结构硬约束】\n"
            f"1. 每章字数必须达到规定区间（{mode} 模式）。\n"
            "2. 每章仅使用 2-3 个 ### 小标题（禁止 4 个及以上，禁止出现孤立的 1 小标题 1 小段）。\n"
            "3. 每个 ### 小标题下必须 ≥3 个自然段、累计 ≥500 字；小标题之间总字数应均衡。\n"
            "4. 每个自然段须按「主张 → 论据（KB 或行业资料）→ 数据 / 表格支撑 → 对项目的含义」四步推进，"
            "避免只写『是什么』而缺少『为什么、多少、意味着什么』。\n"
            "5. 每章必须包含：\n"
            "   ① ≥3 个数字化指标（市场规模 / 增速 / 转化率 / 客单价 / 成本结构等，"
            "若项目方未提供可用行业通用基准，并注明『行业典型水平』）；\n"
            "   ② 1 个 Markdown 表格（竞品对比 / 用户分层 / 收入拆分 / 里程碑 / 风险分级 等，3-5 行 × 2-4 列，"
            "表头用 | 分隔）；\n"
            "   ③ 明确显式使用 1 个分析框架（SWOT / 波特五力 / AARRR / 商业模式画布 / JTBD / 4P / STP / "
            "TAM·SAM·SOM / PEST / 成本结构分层），并在行文中点名该框架。\n\n"
            "【写作风格硬约束】\n"
            "A. 必须利用【行业参考资料】里的事实和数据自然融入行文，不要写『根据搜索结果』或罗列链接。\n"
            "B. 第三人称书面语，严禁『学生 / 用户说 / 在对话中』等元叙述，严禁引用对话原话。\n"
            "C. 语气自信克制，不要反复『尚需验证 / 暂未明确』；整章最多最后一段点一次下一步。\n"
            "D. 财务章节必须结合 KB 中的 budget_facts 给出对商业模式财务合理性的明确判断。\n\n"
            "【反面示例（绝不允许）】\n"
            "> ### 市场规模与趋势\n"
            "> 我国低空经济市场规模巨大，预计到 2035 年将达到 3.5 万亿元。\n"
            "> ### 竞争格局\n"
            "> 目前主要竞品有 A、B、C 三家，各有优劣。\n"
            "（问题：每个小标题只有 1 句话；没有数据铺陈、没有论证、没有对项目的含义）\n\n"
            "【正面示例（每个小标题应达到的饱和度）】\n"
            "> ### 市场规模与趋势（≥500 字，3-4 段）\n"
            "> 第 1 段（主张+宏观数据）：低空经济正从政策引导进入商业落地窗口期，"
            "预计到 2035 年形成 3.5 万亿元规模，年复合增速约 18%……\n"
            "> 第 2 段（结构拆分+行业证据）：按 TAM·SAM·SOM 拆解……\n"
            "> 第 3 段（对项目的含义）：对本项目而言，SAM 约 XXX 亿元，意味着……\n\n"
            "【返回格式（严格）】\n"
            "返回 JSON：{\"sections\": [{\"section_id\", \"content\", \"is_ai_stub\", \"missing_points\"}]}。\n"
            "返回数组 length 必须等于传入 section_id 数量；如某章素材实在匮乏，"
            "也要返回该 sid 并输出详细的框架骨架（不得返回『无法生成』或空 content）。"
        )

    def _upgrade_group(
        self,
        group_ids: list[str],
        kb: dict[str, Any],
        draft_map: dict[str, dict[str, Any]],
        mode: str,
        web_ctx: dict[str, str] | None = None,
        case_header: str = "",
        case_per_section: dict[str, str] | None = None,
    ) -> dict[str, dict[str, Any]]:
        if not self.llm or not self.llm.enabled:
            return {}
        web_ctx = web_ctx or {}
        case_per_section = case_per_section or {}
        specs: list[str] = []
        draft_blocks: list[str] = []
        web_blocks: list[str] = []
        case_blocks: list[str] = []
        for sid in group_ids:
            built = self._build_upgrade_spec(sid, draft_map, mode)
            if not built:
                continue
            spec, draft_block = built
            specs.append(spec)
            draft_blocks.append(draft_block)
            if web_ctx.get(sid):
                web_blocks.append(f"### [{sid} 行业参考资料]\n{web_ctx[sid]}")
            if case_per_section.get(sid):
                case_blocks.append(f"### [{sid} 范本片段]\n{case_per_section[sid]}")

        kb_json = json.dumps(kb, ensure_ascii=False)[:5000]
        system_prompt = self._upgrade_system_prompt(mode)
        if case_header:
            system_prompt += (
                "\n\n" + case_header
                + "\n注意：范本仅供写作风格/深度/结构参考，写作中不得照搬范本的事实、数字、机构名。"
            )
        user_prompt = (
            "【本组章节规格】\n" + "\n".join(specs)
            + "\n\n【项目知识库（KB）— 事实来源】\n" + kb_json
            + ("\n\n【行业参考资料（请自然融入，不要罗列）】\n" + "\n\n".join(web_blocks) if web_blocks else "")
            + ("\n\n【同行业范本片段（仅参考深度与结构）】\n" + "\n\n".join(case_blocks) if case_blocks else "")
            + "\n\n【现有草稿（请在此基础上扩写，不得删除事实点）】\n" + "\n\n".join(draft_blocks)
            + f"\n\n请严格产出 JSON，mode={mode}，sections 长度={len(group_ids)}。"
        )
        try:
            raw = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.4)
        except Exception as exc:
            logger.warning("_upgrade_group failed for %s: %s", group_ids, exc)
            return {}
        if not isinstance(raw, dict):
            return {}
        result: dict[str, dict[str, Any]] = {}
        for item in raw.get("sections") or []:
            if not isinstance(item, dict):
                continue
            sid = str(item.get("section_id") or "").strip()
            content = str(item.get("content") or "").strip()
            if not sid or not content or sid not in group_ids:
                continue
            result[sid] = {
                "content": content,
                "is_ai_stub": bool(item.get("is_ai_stub")),
                "missing_points": item.get("missing_points") if isinstance(item.get("missing_points"), list) else [],
            }
        return result

    def _upgrade_single(
        self,
        sid: str,
        kb: dict[str, Any],
        draft_map: dict[str, dict[str, Any]],
        mode: str,
        web_ctx_text: str = "",
        case_header: str = "",
        case_block_text: str = "",
    ) -> dict[str, Any] | None:
        """单章兜底：只针对一章的窄 prompt。"""
        if not self.llm or not self.llm.enabled:
            return None
        built = self._build_upgrade_spec(sid, draft_map, mode)
        if not built:
            return None
        spec, draft_block = built
        kb_json = json.dumps(kb, ensure_ascii=False)[:4000]
        system_prompt = self._upgrade_system_prompt(mode) + (
            "\n\n【本次特别说明】本次只生成单章，sections 数组长度必须为 1。"
        )
        if case_header:
            system_prompt += (
                "\n\n" + case_header
                + "\n注意：范本仅供写作风格/深度/结构参考，不得照搬事实。"
            )
        user_prompt = (
            "【本章规格】\n" + spec
            + "\n\n【项目知识库（KB）】\n" + kb_json
            + ("\n\n【行业参考资料】\n" + web_ctx_text if web_ctx_text else "")
            + ("\n\n【同行业范本片段】\n" + case_block_text if case_block_text else "")
            + "\n\n【现有草稿】\n" + draft_block
            + f"\n\n请严格产出 JSON：{{\"sections\": [{{ section_id, content, is_ai_stub, missing_points }}]}}，mode={mode}。"
        )
        try:
            raw = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.35)
        except Exception as exc:
            logger.warning("_upgrade_single LLM failed for %s: %s", sid, exc)
            return None
        if not isinstance(raw, dict):
            return None
        for item in raw.get("sections") or []:
            if not isinstance(item, dict):
                continue
            got_sid = str(item.get("section_id") or "").strip()
            content = str(item.get("content") or "").strip()
            if got_sid == sid and content:
                return {
                    "content": content,
                    "is_ai_stub": bool(item.get("is_ai_stub")),
                    "missing_points": item.get("missing_points") if isinstance(item.get("missing_points"), list) else [],
                }
        return None

    # ══════════════════════════════════════════════════════════════
    #  Two-step chapter writer: outline → parallel subhead expand
    # ══════════════════════════════════════════════════════════════

    def _build_chapter_outline(
        self,
        sid: str,
        draft_map: dict[str, dict[str, Any]],
        kb: dict[str, Any],
        web_ctx_text: str = "",
        case_header: str = "",
        case_block_text: str = "",
    ) -> dict[str, Any] | None:
        """
        为一章生成小标题大纲 + 论点骨架。
        返回 {subheads: [{title, thesis, evidence_points[], metrics_hint[], framework, need_table(bool)}]}
        """
        if not self.llm or not self.llm.enabled:
            return None
        template = next((t for t in SECTION_TEMPLATES if t["section_id"] == sid), None)
        if not template:
            return None

        draft = draft_map.get(sid, {})
        draft_content = (draft.get("content") or "")[:1600]
        expected_subheads = (template.get("subheadings") or [])[:3] or ["主要论点", "支撑分析"]
        frameworks = template.get("frameworks") or []
        writing_points = template.get("writing_points") or []

        system_prompt = (
            "你是一位资深商业计划书结构顾问。现在只需要为【一章】规划小标题大纲，不要写正文。\n\n"
            "【硬性要求】\n"
            "1. 严格产出 2-3 个小标题；禁止 4 个及以上。\n"
            "2. 每个小标题附带：\n"
            "   - thesis：一句话核心主张（≥25 字，必须具体、可以被数据支撑）。\n"
            "   - evidence_points：3-4 条论据要点，每条 15-40 字，覆盖『事实 / 数据 / 机制 / 对项目的含义』。\n"
            "   - metrics_hint：2-3 条建议在该小标题下落实的数字化指标或数据坐标。\n"
            "   - framework：该小标题建议使用的一个分析框架名（SWOT/波特五力/AARRR/商业模式画布/"
            "JTBD/4P/STP/TAM·SAM·SOM/PEST/成本结构分层 之一）。\n"
            "   - need_table：true/false，表示该小标题下是否应放一张 Markdown 表格；"
            "整章 2-3 个小标题中只有 1 个 need_table=true。\n"
            "3. thesis 和 evidence_points 必须基于【KB / 行业资料 / 草稿】，不得凭空捏造机构名与数字；"
            "缺乏数据时使用『行业典型水平』语言。\n\n"
            "【返回格式】严格返回 JSON：\n"
            "{\"subheads\": [{\"title\", \"thesis\", \"evidence_points\", \"metrics_hint\", \"framework\", \"need_table\"}]}"
        )
        if case_header:
            system_prompt += "\n\n" + case_header

        kb_json = json.dumps(kb, ensure_ascii=False)[:3500]
        user_prompt = (
            f"章节编号：{sid}\n章节标题：{template['title']}\n"
            f"写作要点参考：{'；'.join(writing_points)}\n"
            f"可选小标题建议（可调整）：{expected_subheads}\n"
            f"可选分析框架：{frameworks}\n\n"
            f"【项目 KB（事实源头）】\n{kb_json}\n\n"
            f"【行业参考资料】\n{web_ctx_text or '（无）'}\n\n"
            f"{('【同行业范本片段（只看结构，不要照搬事实）】\n' + case_block_text) if case_block_text else ''}\n\n"
            f"【现有草稿】\n{draft_content or '（暂无）'}\n\n"
            "请严格输出 JSON，仅包含 subheads 数组，共 2-3 项。"
        )
        try:
            raw = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.3)
        except Exception as exc:
            logger.warning("_build_chapter_outline LLM failed for %s: %s", sid, exc)
            return None
        if not isinstance(raw, dict):
            return None
        subs = raw.get("subheads")
        if not isinstance(subs, list) or not subs:
            return None
        normalized: list[dict[str, Any]] = []
        for item in subs[:3]:
            if not isinstance(item, dict):
                continue
            title = str(item.get("title") or "").strip()
            thesis = str(item.get("thesis") or "").strip()
            if not title or not thesis:
                continue
            ep_raw = item.get("evidence_points") or []
            mh_raw = item.get("metrics_hint") or []
            evidence = [str(x).strip() for x in ep_raw if str(x).strip()][:6]
            metrics = [str(x).strip() for x in mh_raw if str(x).strip()][:4]
            normalized.append({
                "title": title,
                "thesis": thesis,
                "evidence_points": evidence or ["结合 KB 中可得信息推导"],
                "metrics_hint": metrics,
                "framework": str(item.get("framework") or "").strip(),
                "need_table": bool(item.get("need_table")),
            })
        if not normalized:
            return None
        # 保证至少 1 个小标题带表格
        if not any(s["need_table"] for s in normalized):
            normalized[0]["need_table"] = True
        return {"subheads": normalized[:3]}

    def _expand_subhead(
        self,
        sid: str,
        chapter_title: str,
        subhead: dict[str, Any],
        kb: dict[str, Any],
        web_ctx_text: str = "",
        case_block_text: str = "",
        mode: str = "full",
    ) -> str:
        """
        扩写单个小标题为 600-900 字的正文（3-4 段，四要素推进），返回不含 `### 标题` 的正文。
        失败返回用 thesis + evidence 组成的简版。
        """
        title = str(subhead.get("title") or "")
        thesis = str(subhead.get("thesis") or "")
        evidence = list(subhead.get("evidence_points") or [])
        metrics_hint = list(subhead.get("metrics_hint") or [])
        framework = str(subhead.get("framework") or "")
        need_table = bool(subhead.get("need_table"))

        target_range = "600-900" if mode == "full" else "450-700"
        paras = "3-4 段" if mode == "full" else "3 段"

        if not self.llm or not self.llm.enabled:
            return self._render_subhead_fallback(subhead)

        system_prompt = (
            "你是商业计划书撰写顾问。现在只为【一个小标题】写正文，不要再加 ### 标题行。\n\n"
            "【硬性约束】\n"
            f"1. 目标字数 {target_range} 字，共 {paras}；严禁 1-2 句话交差。\n"
            "2. 每段按「主张 → 论据（KB 或行业资料）→ 数据或数字 → 对项目的含义」四步展开；"
            "论据须来自下方 KB / 行业资料 / 范本，不得凭空编造机构名和数字，缺乏数据时使用『行业典型水平』。\n"
            f"3. 必须在行文中明确点名使用分析框架：{framework or '请从适合本小标题的常见框架中选择'}。\n"
            f"4. 必须落实以下指标建议：{('、'.join(metrics_hint)) or '至少给出 2-3 个具体数字'}。\n"
            + ("5. 必须包含 1 个 Markdown 表格（3-5 行 × 2-4 列，表头用 | 分隔），表格承载主要对比或拆解。\n" if need_table else "")
            + "6. 第三人称书面语，严禁『学生 / 用户说 / 在对话中』等元叙述；禁止写 "
            "『根据搜索结果』『以下是我的分析』等口水。\n\n"
            "【返回】纯 Markdown 正文（不要包 ### 标题，不要 JSON，不要代码块包裹），可含段落、列表、表格。"
        )
        kb_json = json.dumps(kb, ensure_ascii=False)[:2600]
        evidence_block = "\n".join(f"- {e}" for e in evidence)
        user_prompt = (
            f"所属章节：{chapter_title}（{sid}）\n"
            f"小标题：{title}\n"
            f"核心主张：{thesis}\n"
            f"必须展开的论据点：\n{evidence_block}\n\n"
            f"【项目 KB】\n{kb_json}\n\n"
            f"【行业参考资料】\n{web_ctx_text or '（无）'}\n\n"
            f"{('【同行业范本片段（只借鉴结构，不照抄事实）】\n' + case_block_text) if case_block_text else ''}\n\n"
            "现在直接输出该小标题的正文 Markdown（不带 ### 标题行）。"
        )
        try:
            text = self.llm.chat_text(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                temperature=0.4,
            )
        except Exception as exc:
            logger.warning("_expand_subhead LLM failed for %s/%s: %s", sid, title, exc)
            return self._render_subhead_fallback(subhead)

        if not text:
            return self._render_subhead_fallback(subhead)
        cleaned = str(text).strip()
        # 去掉模型自作主张加上的 ### 标题头
        if cleaned.startswith("### "):
            lines = cleaned.splitlines()
            cleaned = "\n".join(lines[1:]).strip()
        # 去掉 ```markdown 包裹
        if cleaned.startswith("```"):
            parts = cleaned.split("```")
            if len(parts) >= 2:
                cleaned = parts[1]
                if cleaned.lower().startswith("markdown"):
                    cleaned = cleaned[len("markdown"):]
                cleaned = cleaned.strip()
        return cleaned or self._render_subhead_fallback(subhead)

    def _render_subhead_fallback(self, subhead: dict[str, Any]) -> str:
        """单小标题扩写失败时的兜底：用 outline 的 thesis + evidence 输出一段可读文本。"""
        thesis = str(subhead.get("thesis") or "").strip()
        evidence = list(subhead.get("evidence_points") or [])
        metrics = list(subhead.get("metrics_hint") or [])
        framework = str(subhead.get("framework") or "").strip()
        parts: list[str] = []
        if thesis:
            parts.append(thesis)
        if evidence:
            parts.append("本节关键支撑要点如下：")
            parts.append("\n".join(f"- {e}" for e in evidence))
        if metrics:
            parts.append("建议在后续细化中落实的数字化指标：" + "、".join(metrics) + "。")
        if framework:
            parts.append(f"分析框架：{framework}。")
        parts.append("（AI 扩写此小标题时遇到临时问题，以上为大纲兜底，建议再次点击升级。）")
        return "\n\n".join(parts)

    def _upgrade_chapter_two_step(
        self,
        sid: str,
        kb: dict[str, Any],
        draft_map: dict[str, dict[str, Any]],
        mode: str,
        web_ctx_text: str = "",
        case_header: str = "",
        case_block_text: str = "",
    ) -> dict[str, Any] | None:
        """
        整章两步式：先拿 outline，再并发扩写每个小标题，最后拼装成 Markdown content。
        """
        outline = self._build_chapter_outline(
            sid, draft_map, kb,
            web_ctx_text=web_ctx_text,
            case_header=case_header,
            case_block_text=case_block_text,
        )
        if not outline or not outline.get("subheads"):
            return None
        subheads = list(outline["subheads"])
        template = next((t for t in SECTION_TEMPLATES if t["section_id"] == sid), None)
        chapter_title = (template or {}).get("title") or sid

        expanded: dict[int, str] = {}
        # 并发每个小标题
        with ThreadPoolExecutor(max_workers=min(3, len(subheads))) as pool:
            futures = {
                pool.submit(
                    self._expand_subhead,
                    sid, chapter_title, sh,
                    kb, web_ctx_text, case_block_text, mode,
                ): idx
                for idx, sh in enumerate(subheads)
            }
            for fut, idx in futures.items():
                try:
                    expanded[idx] = fut.result(timeout=120)
                except Exception as exc:
                    logger.warning("subhead expand failed: %s (%s) %s", sid, subheads[idx].get("title"), exc)
                    expanded[idx] = self._render_subhead_fallback(subheads[idx])

        blocks: list[str] = []
        for idx, sh in enumerate(subheads):
            body = (expanded.get(idx) or "").strip() or self._render_subhead_fallback(sh)
            blocks.append(f"### {sh['title']}\n\n{body}")
        content = "\n\n".join(blocks).strip()
        if not content:
            return None
        # 校验最小字数
        total_chinese = self._chinese_length(content)
        is_stub = total_chinese < (900 if mode == "basic" else 1500)
        return {
            "content": content,
            "is_ai_stub": is_stub,
            "missing_points": [],
            "outline": outline,
        }

    # ══════════════════════════════════════════════════════════════
    #  Batch revision accept / reject
    # ══════════════════════════════════════════════════════════════

    def accept_all_revisions(self, plan_id: str) -> dict[str, Any]:
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found", "plan": None, "accepted": 0}
        revisions = list(plan.get("pending_revisions") or [])
        if not revisions:
            return {"status": "ok", "plan": plan, "accepted": 0}
        previous = copy.deepcopy(plan)
        plan_copy = copy.deepcopy(plan)
        sections = list(plan_copy.get("sections") or [])
        sec_by_id = {str(s.get("section_id")): s for s in sections}
        accepted = 0
        for rev in revisions:
            sid = str(rev.get("section_id") or "")
            new_content = str(rev.get("new_content") or rev.get("ai_draft") or "").strip()
            if not sid or not new_content:
                continue
            section = sec_by_id.get(sid)
            if not section:
                continue
            history = list(section.get("content_history") or [])
            history.append({
                "content": section.get("content") or "",
                "replaced_by": f"accept:{rev.get('revision_id')}",
                "replaced_at": _now_iso(),
            })
            section["content_history"] = history[-3:]
            section["content"] = new_content
            section["user_edit"] = ""
            section["ai_draft"] = ""
            section["revision_status"] = "accepted"
            if rev.get("candidate_missing_points"):
                section["missing_points"] = rev["candidate_missing_points"]
            if rev.get("candidate_missing_level"):
                section["missing_level"] = rev["candidate_missing_level"]
            accepted += 1
        plan_copy["sections"] = sections
        plan_copy["pending_revisions"] = []
        plan_copy["revision_badge_count"] = 0
        plan_copy["status"] = "accepted" if accepted else plan_copy.get("status", "draft")
        plan_copy["updated_at"] = _now_iso()
        saved = self._save_plan(plan_copy, previous=previous)
        return {"status": "ok", "plan": saved, "accepted": accepted}

    def reject_all_revisions(self, plan_id: str) -> dict[str, Any]:
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found", "plan": None, "rejected": 0}
        revisions = list(plan.get("pending_revisions") or [])
        if not revisions:
            return {"status": "ok", "plan": plan, "rejected": 0}
        previous = copy.deepcopy(plan)
        plan_copy = copy.deepcopy(plan)
        sections = list(plan_copy.get("sections") or [])
        for section in sections:
            if section.get("ai_draft"):
                section["ai_draft"] = ""
            if section.get("revision_status") == "pending":
                section["revision_status"] = "rejected"
        plan_copy["sections"] = sections
        rejected = len(revisions)
        plan_copy["pending_revisions"] = []
        plan_copy["revision_badge_count"] = 0
        plan_copy["updated_at"] = _now_iso()
        saved = self._save_plan(plan_copy, previous=previous)
        return {"status": "ok", "plan": saved, "rejected": rejected}

    # ══════════════════════════════════════════════════════════════
    #  Deepen Questions + Expand Single Section
    # ══════════════════════════════════════════════════════════════

    def generate_deepen_questions(self, plan_id: str, section_id: str) -> dict[str, Any]:
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found", "questions": []}
        section = next((s for s in plan.get("sections") or [] if str(s.get("section_id")) == section_id), None)
        if not section:
            return {"status": "not_found", "questions": []}
        kb = plan.get("knowledge_base") or {}

        if not self.llm or not self.llm.enabled:
            return {"status": "ok", "questions": self._fallback_deepen_questions(section, kb)}

        kb_json = json.dumps(kb, ensure_ascii=False)[:3500]
        section_preview = str(section.get("content") or "")[:1500]
        system_prompt = (
            "你是一位资深商业计划书顾问，现在要针对用户的某一章节，提出 3-5 个具体、可回答的深化问题，"
            "帮助用户把这一章写得更扎实。硬性要求：\n"
            "1. 问题围绕该章的写作要点与常用分析框架展开，并结合 KB 里的事实；不要重复草稿里已经写过的内容。\n"
            "2. 每个问题都应是 '场景/数据/决策理由' 层面的具体问题，而不是 '你们的目标用户是谁' 这种宽泛问题。\n"
            "3. 用第三人称客观口吻提问，不要出现 '你/您/同学'。\n"
            "返回 JSON：{\"questions\": [{\"id\", \"text\", \"focus_point\"}]}，id 用字母+数字如 q1/q2。"
        )
        user_prompt = (
            f"【章节】{section.get('display_title') or section.get('title')}\n"
            f"【章节正文（草稿）】\n{section_preview}\n\n"
            f"【项目知识库 KB】\n{kb_json}\n\n"
            "请返回 3-5 个深化问题的 JSON。"
        )
        try:
            raw = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.4)
        except Exception as exc:
            logger.warning("generate_deepen_questions failed: %s", exc)
            return {"status": "ok", "questions": self._fallback_deepen_questions(section, kb)}
        questions = []
        for idx, item in enumerate((raw or {}).get("questions") or []):
            if not isinstance(item, dict):
                continue
            text = str(item.get("text") or "").strip()
            if not text:
                continue
            questions.append({
                "id": str(item.get("id") or f"q{idx + 1}"),
                "text": text,
                "focus_point": str(item.get("focus_point") or "").strip(),
            })
        if not questions:
            questions = self._fallback_deepen_questions(section, kb)
        return {"status": "ok", "questions": questions}

    def _fallback_deepen_questions(self, section: dict[str, Any], kb: dict[str, Any]) -> list[dict[str, Any]]:
        title = section.get("display_title") or section.get("title") or "本章"
        return [
            {"id": "q1", "text": f"{title} 中最关键的 1-2 条事实或数字是什么？", "focus_point": "事实与数据"},
            {"id": "q2", "text": f"针对 {title}，已经做过哪些验证？结论如何？", "focus_point": "验证与证据"},
            {"id": "q3", "text": f"如果外部环境变化，{title} 里最脆弱的假设是什么？", "focus_point": "风险与敏感度"},
        ]

    def expand_section(
        self,
        plan_id: str,
        section_id: str,
        *,
        answers: list[dict[str, Any]],
        merge_strategy: str = "append",
    ) -> dict[str, Any]:
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found", "plan": None}
        section = next((s for s in plan.get("sections") or [] if str(s.get("section_id")) == section_id), None)
        if not section:
            return {"status": "not_found", "plan": None}

        kb = plan.get("knowledge_base") or {}
        template = next((t for t in SECTION_TEMPLATES if str(t["section_id"]) == section_id), None)
        original_content = str(section.get("content") or "")
        new_content = ""
        if self.llm and self.llm.enabled and template:
            frameworks = "、".join(template.get("frameworks") or [])
            subheads = "、".join(template.get("subheadings") or [])
            answers_block = "\n".join(
                f"- Q: {str(a.get('question_text') or a.get('question_id') or '').strip()}\n  A: {str(a.get('text') or '').strip()}"
                for a in (answers or [])
                if isinstance(a, dict) and str(a.get("text") or "").strip()
            ) or "（未提供具体答复）"
            kb_json = json.dumps(kb, ensure_ascii=False)[:4500]
            system_prompt = (
                "你是一位资深商业计划书顾问，现在要把某一章扩写到 2500-4000 字的正式版本，"
                "融合学生对深化问题的答复、KB 中的事实、以及本章应使用的分析框架。\n"
                "【硬性要求】\n"
                "1. 第三人称书面语，禁止 '学生 / 用户说 / 在对话中' 等元叙述，也禁止引用原话。\n"
                "2. 必须使用以下分析框架之一或组合：" + (frameworks or "行业通用框架") + "。\n"
                "3. 至少包含 4-6 段、3-4 个 ### 小标题；段与段之间有论点-论证-推论推进。\n"
                "4. 如果现有章节已经有事实点，必须保留，不得删减；答复中的新信息自然融入行文。\n"
                "5. 默认语气自信克制，整章最多在最后一段点一次下一步。\n"
                "返回 JSON：{\"content\": \"Markdown 正文\"}。"
            )
            user_prompt = (
                f"【章节】{template.get('title')}\n"
                f"【建议小标题】{subheads}\n"
                f"【现有章节正文】\n{original_content}\n\n"
                f"【用户对深化问题的答复】\n{answers_block}\n\n"
                f"【项目知识库 KB】\n{kb_json}\n\n"
                "请返回扩写后的完整 Markdown 正文的 JSON。"
            )
            try:
                raw = self.llm.chat_json(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.4)
                new_content = str((raw or {}).get("content") or "").strip()
            except Exception as exc:
                logger.warning("expand_section failed: %s", exc)

        if not new_content:
            return {"status": "llm_unavailable", "plan": None}

        if merge_strategy == "append" and original_content and original_content not in new_content:
            new_content = original_content.rstrip() + "\n\n---\n\n" + new_content

        previous = copy.deepcopy(plan)
        next_sections: list[dict[str, Any]] = []
        revision = None
        for row in plan.get("sections") or []:
            row_copy = copy.deepcopy(row)
            if str(row_copy.get("section_id")) == section_id:
                revision = {
                    "revision_id": str(uuid4())[:8],
                    "section_id": section_id,
                    "section_title": row_copy.get("display_title") or row_copy.get("title"),
                    "summary": f"{row_copy.get('display_title') or row_copy.get('title')} 已深化扩写",
                    "reason": "基于深化问答的新补充内容，已合并到本章。",
                    "source_hint": "深化本章",
                    "old_content": row_copy.get("content") or "",
                    "new_content": new_content,
                    "candidate_field_map": row_copy.get("field_map") or {},
                    "candidate_missing_points": row_copy.get("missing_points") or [],
                    "candidate_missing_level": "complete",
                    "changes": self._build_inline_changes(row_copy.get("content") or "", new_content),
                    "created_at": _now_iso(),
                }
                row_copy["ai_draft"] = new_content
                row_copy["revision_status"] = "pending"
                row_copy["is_ai_stub"] = False
            next_sections.append(row_copy)

        plan_copy = copy.deepcopy(plan)
        plan_copy["sections"] = next_sections
        pending = list(plan_copy.get("pending_revisions") or [])
        if revision:
            pending.append(revision)
        plan_copy["pending_revisions"] = pending
        plan_copy["revision_badge_count"] = len(pending)
        plan_copy["updated_at"] = _now_iso()
        saved = self._save_plan(plan_copy, previous=previous)
        return {"status": "ok", "plan": saved, "revision": revision}

    # ══════════════════════════════════════════════════════════════
    #  Per-chapter deepen loop：3-5 个靶向问题 → 用户回答 → 改写该章
    # ══════════════════════════════════════════════════════════════

    def generate_chapter_deepen_questions(
        self, plan_id: str, section_id: str
    ) -> dict[str, Any]:
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found", "questions": []}
        sections = list(plan.get("sections") or [])
        section = next((s for s in sections if str(s.get("section_id")) == section_id), None)
        if not section:
            return {"status": "not_found", "questions": []}

        title = section.get("display_title") or section.get("title") or section_id
        kb = plan.get("knowledge_base") or {}
        content_preview = str(section.get("content") or "")[:1800]
        missing_points = section.get("missing_points") or []
        template = next((t for t in SECTION_TEMPLATES if t["section_id"] == section_id), None)
        writing_points = (template or {}).get("writing_points") or []

        fallback_questions = [
            {"id": f"q{i+1}",
             "question": f"关于「{title}」，请补充第 {i+1} 个关键事实或具体数字？",
             "why": "该章节尚缺关键细节", "hint": ""}
            for i in range(3)
        ]

        if not self.llm or not self.llm.enabled:
            return {"status": "ok", "questions": fallback_questions}

        system_prompt = (
            "你是商业计划书顾问，需要针对指定章节设计 3-5 个靶向问题，"
            "帮助学生团队把该章写得更深入、更具体。\n\n"
            "【要求】\n"
            "1. 每个问题必须针对该章的关键信息空缺，避免空泛的『你如何看待...』。\n"
            "2. 问题必须具体、可回答，形式如『请给出...的具体数字或案例』、『...的核心假设是什么』。\n"
            "3. 问题之间覆盖不同维度（数据 / 流程 / 案例 / 风险 / 差异化），避免重复。\n"
            "4. 禁止使用『你 / 您 / 团队你们』等第二人称，一律第三人称书面语。\n"
            "5. 每个问题附 why（为什么需要这个信息）与 hint（可能的回答线索）。\n\n"
            "返回严格 JSON：\n"
            "{\"questions\": [{\"id\":\"q1\",\"question\":\"...\",\"why\":\"...\",\"hint\":\"...\"}]}，"
            "数组长度 3-5。"
        )
        user_prompt = (
            f"章节编号：{section_id}\n"
            f"章节标题：{title}\n"
            f"本章写作要点：{'；'.join(writing_points) or '（见通用模板）'}\n"
            f"现有 missing_points：{missing_points}\n\n"
            f"【现有章节内容（截断）】\n{content_preview or '（暂无）'}\n\n"
            f"【项目 KB】\n{json.dumps(kb, ensure_ascii=False)[:2400]}\n\n"
            "请输出针对本章的 3-5 个深化问题。"
        )
        try:
            raw = self.llm.chat_json(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                temperature=0.45,
            )
        except Exception as exc:
            logger.warning("generate_chapter_deepen_questions LLM failed: %s", exc)
            return {"status": "ok", "questions": fallback_questions}
        if not isinstance(raw, dict):
            return {"status": "ok", "questions": fallback_questions}
        items = raw.get("questions")
        if not isinstance(items, list) or not items:
            return {"status": "ok", "questions": fallback_questions}

        out: list[dict[str, Any]] = []
        for i, item in enumerate(items[:5]):
            if not isinstance(item, dict):
                continue
            q = str(item.get("question") or "").strip()
            if not q:
                continue
            out.append({
                "id": str(item.get("id") or f"q{i+1}"),
                "question": q,
                "why": str(item.get("why") or "").strip(),
                "hint": str(item.get("hint") or "").strip(),
            })
        if len(out) < 3:
            out = out + fallback_questions[: 3 - len(out)]
        return {"status": "ok", "section_id": section_id, "section_title": title, "questions": out}

    def apply_chapter_deepen(
        self, plan_id: str, section_id: str, answers: list[dict[str, Any]]
    ) -> dict[str, Any]:
        """
        学生提交问题答案 → 写入 KB.chapter_addons[sid] → 针对该章重新跑两步式 → 产出 revision。
        """
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found", "plan": None, "revision": None}
        sections = list(plan.get("sections") or [])
        section = next((s for s in sections if str(s.get("section_id")) == section_id), None)
        if not section:
            return {"status": "not_found", "plan": None, "revision": None}

        # 1. 把 answers 写到 knowledge_base.chapter_addons[sid]（累积）
        kb = copy.deepcopy(plan.get("knowledge_base") or {})
        addons_map = kb.get("chapter_addons") or {}
        if not isinstance(addons_map, dict):
            addons_map = {}
        existing: list[str] = list(addons_map.get(section_id) or [])
        new_lines: list[str] = []
        for ans in answers or []:
            if not isinstance(ans, dict):
                continue
            question = str(ans.get("question") or "").strip()
            answer = str(ans.get("answer") or "").strip()
            if not answer:
                continue
            if question:
                new_lines.append(f"Q：{question}\nA：{answer}")
            else:
                new_lines.append(answer)
        if not new_lines:
            return {
                "status": "no_answers",
                "plan": plan,
                "revision": None,
                "message": "未收到任何有效答案，未触发改写。",
            }
        addons_map[section_id] = existing + new_lines
        kb["chapter_addons"] = addons_map

        # 2. 基于新 KB 对这一章重新跑两步式
        title = section.get("display_title") or section.get("title") or section_id
        draft_map = {
            section_id: {
                "title": title,
                "content": section.get("content") or "",
                "is_ai_stub": bool(section.get("is_ai_stub")),
                "has_material": True,
            }
        }
        web_ctx = self._collect_web_context([section_id], kb)
        case_header, case_per_section = self._retrieve_case_fewshots(kb)
        mode = str(plan.get("version_tier") or "full")
        if mode not in UPGRADE_LEN:
            mode = "full"

        result: dict[str, Any] | None = None
        try:
            result = self._upgrade_chapter_two_step(
                section_id, kb, draft_map, mode,
                web_ctx.get(section_id, ""),
                case_header,
                case_per_section.get(section_id, ""),
            )
        except Exception as exc:
            logger.warning("apply_chapter_deepen two-step failed: %s", exc)
            result = None

        if not result or not result.get("content"):
            try:
                result = self._upgrade_single(
                    section_id, kb, draft_map, mode,
                    web_ctx.get(section_id, ""),
                    case_header, case_per_section.get(section_id, ""),
                )
            except Exception as exc:
                logger.warning("apply_chapter_deepen single fallback failed: %s", exc)

        if not result or not result.get("content"):
            return {
                "status": "error",
                "plan": plan,
                "revision": None,
                "message": "改写失败：LLM 临时不可用，请稍后再试。",
            }

        new_content = str(result.get("content")).strip()
        old_content = str(section.get("content") or "")

        # 3. 生成 revision 写回 plan
        previous = copy.deepcopy(plan)
        plan_copy = copy.deepcopy(plan)
        plan_copy["knowledge_base"] = kb

        revision = None
        next_sections: list[dict[str, Any]] = []
        for s in plan_copy.get("sections") or []:
            row = copy.deepcopy(s)
            if str(row.get("section_id")) == section_id and new_content and new_content != old_content:
                revision = {
                    "revision_id": str(uuid4())[:8],
                    "section_id": section_id,
                    "section_title": title,
                    "summary": f"{title} 基于深化答题重写",
                    "reason": "基于学生针对本章的补充答题整合进 KB，并按两步式流水线重写本章。",
                    "source_hint": "chapter_deepen",
                    "old_content": old_content,
                    "new_content": new_content,
                    "candidate_field_map": row.get("field_map") or {},
                    "candidate_missing_points": result.get("missing_points") or [],
                    "candidate_missing_level": "complete",
                    "changes": self._build_inline_changes(old_content, new_content),
                    "created_at": _now_iso(),
                }
                row["ai_draft"] = new_content
                row["revision_status"] = "pending"
                if "is_ai_stub" in result:
                    row["is_ai_stub"] = bool(result["is_ai_stub"])
                if result.get("outline"):
                    row["outline"] = result["outline"]
            next_sections.append(row)

        plan_copy["sections"] = next_sections
        pending = list(plan_copy.get("pending_revisions") or [])
        if revision:
            pending.append(revision)
        plan_copy["pending_revisions"] = pending
        plan_copy["revision_badge_count"] = len(pending)
        plan_copy["updated_at"] = _now_iso()
        saved = self._save_plan(plan_copy, previous=previous)
        return {"status": "ok", "plan": saved, "revision": revision}

    def generate_deepen_suggestions(self, plan_id: str) -> dict[str, Any]:
        plan = self.storage.load(plan_id)
        if not plan:
            return {"status": "not_found", "suggestions": []}
        sections = list(plan.get("sections") or [])
        kb = plan.get("knowledge_base") or {}

        scored: list[tuple[float, dict[str, Any]]] = []
        for section in sections:
            content_len = self._chinese_length(str(section.get("content") or ""))
            missing_level = str(section.get("missing_level") or "")
            is_stub = bool(section.get("is_ai_stub"))
            priority = 0.0
            if missing_level == "critical":
                priority += 3
            elif missing_level == "partial":
                priority += 2
            if is_stub:
                priority += 2
            if content_len < 400:
                priority += 1.5
            priority += len(section.get("missing_points") or []) * 0.3
            scored.append((priority, section))

        scored.sort(key=lambda x: x[0], reverse=True)
        top = [item for item in scored if item[0] > 0][:8]
        suggestions: list[dict[str, Any]] = []
        for priority, section in top:
            sid = str(section.get("section_id"))
            title = section.get("display_title") or section.get("title") or sid
            question = f"{title} 还需要补充哪些关键事实或案例？"
            if self.llm and self.llm.enabled:
                try:
                    raw = self.llm.chat_json(
                        system_prompt=(
                            "你是商业计划书顾问，需要针对某一章提出 1 个最关键的深化问题。"
                            "问题要具体、可回答，避免使用 '你/您'。"
                            "返回 JSON：{\"question\":\"...\",\"why\":\"简短原因\"}"
                        ),
                        user_prompt=(
                            f"章节：{title}\n"
                            f"现有内容（截断）：{str(section.get('content') or '')[:900]}\n"
                            f"KB 摘要：{json.dumps(kb, ensure_ascii=False)[:1500]}"
                        ),
                        temperature=0.5,
                    )
                    if isinstance(raw, dict) and raw.get("question"):
                        question = str(raw.get("question"))
                        why = str(raw.get("why") or "")
                    else:
                        why = ""
                except Exception:
                    why = ""
            else:
                why = "该章节内容较为简略或为 AI 参考稿"
            suggestions.append({
                "section_id": sid,
                "section_title": title,
                "priority": round(priority, 2),
                "question": question,
                "why": why,
            })

        return {"status": "ok", "suggestions": suggestions}

    # ══════════════════════════════════════════════════════════════
    #  Budget
    # ══════════════════════════════════════════════════════════════

    def _load_budget_hint(self, student_id: str) -> dict[str, Any]:
        if not student_id:
            return {}
        budget_root = self.json_store.root.parent / "budgets" / student_id
        if not budget_root.exists():
            return {}
        files = sorted(budget_root.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
        for file in files:
            try:
                data = json.loads(file.read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    return data
            except Exception:
                continue
        return {}
